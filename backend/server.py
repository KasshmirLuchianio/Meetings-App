import os
import json
import uuid
import re
import math
import glob
import shutil
import asyncio
import secrets
import tempfile
from datetime import datetime, timezone, timedelta
from pathlib import Path
from io import BytesIO
from typing import Optional, List

from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Query, Request, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, HTMLResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, Field
from dotenv import load_dotenv
from motor.motor_asyncio import AsyncIOMotorClient
from bson import ObjectId
import base64
import aiofiles
import httpx
import stripe
import jwt
import bcrypt
import resend
from slowapi import Limiter
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from openai import AsyncOpenAI
from anthropic import AsyncAnthropic

load_dotenv()

# ==================== CONFIG ====================
MONGO_URL = os.environ.get("MONGO_URL", "mongodb://localhost:27017/gal_meetings")
DB_NAME = os.environ.get("DB_NAME", "gal_meetings")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")

# ==================== STRIPE ====================
stripe.api_key = os.environ.get("STRIPE_SECRET_KEY")
STRIPE_WEBHOOK_SECRET = os.environ.get("STRIPE_WEBHOOK_SECRET")
STRIPE_PRICES = {
    "pro_monthly": "price_1TIOnGHa4KY3ww8wNuZfnCvN",
    "pro_yearly": "price_1TIOneHa4KY3ww8wkmyVhjDP",
    "enterprise_monthly": "price_1TIOnuHa4KY3ww8wYMLsIRIw",
    "enterprise_yearly": "price_1TIOoCHa4KY3ww8wjDYteZ14",
}

# ==================== RESEND (Email) ====================
resend.api_key = os.environ.get("RESEND_API_KEY")

# ==================== SMARTBILL ====================
SMARTBILL_EMAIL = "vladgrigorov1@gmail.com"
SMARTBILL_API_KEY = os.environ.get("SMARTBILL_API_KEY", "")
SMARTBILL_CUI = "46076724"
SMARTBILL_SERIES = "Meetings"
SMARTBILL_BASE_URL = "https://ws.smartbill.ro/SBORO/api"

EUR_TO_RON = 4.97
PLAN_RON_PRICES = {
    "pro": round(19 * EUR_TO_RON, 2),         # 94.43 RON
    "enterprise": round(99 * EUR_TO_RON, 2),  # 491.03 RON
}
PLAN_EUR_PRICES = {"pro": 19.0, "enterprise": 99.0}


def _smartbill_auth_header() -> str:
    """Basic auth header for SmartBill SBORO API."""
    token = base64.b64encode(f"{SMARTBILL_EMAIL}:{SMARTBILL_API_KEY}".encode()).decode()
    return f"Basic {token}"


async def create_smartbill_invoice(
    client_name: str,
    client_email: str,
    client_cui: Optional[str],
    client_address: Optional[str],
    plan: str,  # "pro" or "enterprise"
    issue_date: str,  # "YYYY-MM-DD"
) -> Optional[str]:
    """
    Create a SmartBill invoice and return the invoice series+number (e.g. "Meetings-001").
    Returns None on failure (non-blocking — payment already succeeded).
    """
    if not SMARTBILL_API_KEY:
        print("[SmartBill] SMARTBILL_API_KEY not configured — skipping invoice creation")
        return None

    plan_key = plan.lower()
    unit_price_ron = PLAN_RON_PRICES.get(plan_key)
    unit_price_eur = PLAN_EUR_PRICES.get(plan_key)
    if unit_price_ron is None:
        print(f"[SmartBill] Unknown plan '{plan}' — skipping")
        return None

    # VAT 19% (Romania standard rate)
    vat_rate = 19.0
    vat_amount = round(unit_price_ron * vat_rate / 100, 2)
    total_ron = round(unit_price_ron + vat_amount, 2)

    product_name = f"Abonament Meetings.ro {plan.capitalize()} (1 lună)"
    invoice_payload = {
        "companyVatCode": SMARTBILL_CUI,
        "client": {
            "name": client_name or client_email,
            "vatCode": client_cui or "",
            "address": client_address or "",
            "email": client_email,
            "saveToDb": False,
            "isTaxPayer": bool(client_cui),
        },
        "issueDate": issue_date,
        "seriesName": SMARTBILL_SERIES,
        "isDraft": False,
        "dueDate": issue_date,
        "currency": "RON",
        "language": "RO",
        "precision": 2,
        "products": [
            {
                "name": product_name,
                "measuringUnitName": "buc",
                "currency": "RON",
                "quantity": 1,
                "unitPrice": unit_price_ron,
                "isTaxIncluded": False,
                "taxName": "Normala",
                "taxPercentage": vat_rate,
                "isDiscount": False,
                "saveToDb": False,
            }
        ],
        "observations": (
            f"Servicii SaaS — transcriere și analiză ședințe. "
            f"Preț EUR: {unit_price_eur}€ × curs {EUR_TO_RON} = {unit_price_ron} RON + TVA."
        ),
        "paymentType": "Card",
        "paymentDate": issue_date,
        "useStock": False,
    }

    try:
        async with httpx.AsyncClient(timeout=15.0) as hclient:
            resp = await hclient.post(
                f"{SMARTBILL_BASE_URL}/invoice",
                json=invoice_payload,
                headers={
                    "Authorization": _smartbill_auth_header(),
                    "Content-Type": "application/json",
                    "Accept": "application/json",
                },
            )
        if resp.status_code in (200, 201):
            data = resp.json()
            series = data.get("series", SMARTBILL_SERIES)
            number = data.get("number", "")
            invoice_ref = f"{series}-{number}"
            print(f"[SmartBill] Invoice created: {invoice_ref} for {client_email} ({plan})")
            return invoice_ref
        else:
            print(f"[SmartBill] Invoice creation failed {resp.status_code}: {resp.text[:300]}")
            return None
    except Exception as exc:
        print(f"[SmartBill] Exception during invoice creation: {exc}")
        return None


async def send_efactura_to_anaf(invoice_series: str, invoice_number: str) -> bool:
    """
    Send the invoice to ANAF e-Factura system via SmartBill.
    Returns True on success.
    """
    if not SMARTBILL_API_KEY:
        return False
    try:
        async with httpx.AsyncClient(timeout=15.0) as hclient:
            resp = await hclient.post(
                f"{SMARTBILL_BASE_URL}/invoice/efactura",
                json={
                    "companyVatCode": SMARTBILL_CUI,
                    "seriesName": invoice_series,
                    "invoice": invoice_number,
                    "uploadInvoice": True,
                },
                headers={
                    "Authorization": _smartbill_auth_header(),
                    "Content-Type": "application/json",
                    "Accept": "application/json",
                },
            )
        if resp.status_code in (200, 201):
            print(f"[SmartBill] e-Factura sent to ANAF: {invoice_series}-{invoice_number}")
            return True
        else:
            print(f"[SmartBill] e-Factura failed {resp.status_code}: {resp.text[:200]}")
            return False
    except Exception as exc:
        print(f"[SmartBill] e-Factura exception: {exc}")
        return False


async def _create_and_save_invoice(user_email: str, plan: str, issue_date: str) -> None:
    """
    Full invoice flow: look up user billing info → create SmartBill invoice →
    optionally send e-Factura → save invoice record to MongoDB.
    Designed to be called as asyncio.create_task() so it never blocks Stripe webhook response.
    """
    try:
        user = await db.users.find_one({"email": user_email})
        billing_name = (user or {}).get("billing_name", "") or user_email
        billing_cui = (user or {}).get("billing_cui", "")
        billing_address = (user or {}).get("billing_address", "")

        invoice_ref = await create_smartbill_invoice(
            client_name=billing_name,
            client_email=user_email,
            client_cui=billing_cui,
            client_address=billing_address,
            plan=plan,
            issue_date=issue_date,
        )

        if invoice_ref:
            # Optionally push to ANAF e-Factura (best-effort)
            parts = invoice_ref.split("-", 1)
            if len(parts) == 2:
                await send_efactura_to_anaf(parts[0], parts[1])

        # Persist invoice record regardless of SmartBill success
        await db.invoices.insert_one({
            "user_email": user_email,
            "plan": plan,
            "amount_ron": PLAN_RON_PRICES.get(plan.lower()),
            "amount_eur": PLAN_EUR_PRICES.get(plan.lower()),
            "invoice_ref": invoice_ref,
            "issue_date": issue_date,
            "created_at": datetime.now(timezone.utc),
        })
    except Exception as exc:
        print(f"[SmartBill] _create_and_save_invoice error for {user_email}: {exc}")

# CORS Origins - includes Expo development
CORS_ALLOWED_ORIGINS = [
    "http://localhost:3000",                         # Web dev
    "https://*.onrender.com",                         # Render preview
    "https://meetings.ro",                           # Production web
    "exp://192.168.*",                               # Expo Go LAN (iOS/Android)
    "exp://localhost:19000",                         # Expo dev server
    "capacitor://localhost",                         # Capacitor (if needed)
]

UPLOAD_DIR = Path("/app/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def cleanup_orphaned_temp_files():
    """
    TASK 8 — Remove leftover temp WAV files from jobs that crashed mid-pipeline.
    Patterns: *_preprocessed.wav, *_chunk_NNN.wav, *_clean.wav
    Searches both UPLOAD_DIR and the system temp directory.
    Called once at process startup.
    """
    search_dirs = [str(UPLOAD_DIR), tempfile.gettempdir()]
    patterns = ["*_preprocessed.wav", "*_chunk_*.wav", "*_clean.wav"]
    removed = 0
    for directory in search_dirs:
        for pattern in patterns:
            for path in glob.glob(os.path.join(directory, pattern)):
                try:
                    os.remove(path)
                    removed += 1
                    print(f"[Startup] Removed orphan temp file: {path}")
                except Exception as e:
                    print(f"[Startup] Could not remove {path}: {e}")
    if removed:
        print(f"[Startup] Cleaned {removed} orphaned temp file(s)")
    else:
        print("[Startup] No orphaned temp files found")


cleanup_orphaned_temp_files()

# ==================== APP ====================
limiter = Limiter(key_func=get_remote_address)
app = FastAPI(title="Meetings.ro API", version="1.0.0")
app.state.limiter = limiter


@app.exception_handler(RateLimitExceeded)
async def rate_limit_handler(request: Request, exc: RateLimitExceeded):
    return JSONResponse(
        status_code=429,
        content={"detail": "Prea multe cereri. Încearcă din nou mai târziu."}
    )

ALLOWED_ORIGINS = [
    "https://meetings-ro-api.onrender.com",
    "http://localhost:8081",
    "exp://localhost:8081",
    "exp://192.168.*",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Range", "Accept-Ranges"],
)

# ==================== AUTH CONFIG ====================
JWT_SECRET = os.environ.get("JWT_SECRET")
if not JWT_SECRET:
    raise RuntimeError("JWT_SECRET environment variable is required")
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 72  # 3 days

# ==================== PLAN LIMITS ====================
# Minute-based monthly quota. Stored plan strings in DB are uppercase
# ("FREE"/"PRO"/"ENTERPRISE") for legacy reasons — normalize_plan() maps them
# to the lowercase keys used here.
PLAN_LIMITS = {
    "starter":    {"minutes": 15,    "label": "15 minute/lună"},
    "pro":        {"minutes": 300,   "label": "5 ore/lună"},
    "enterprise": {"minutes": 99999, "label": "Nelimitat"},
}


def normalize_plan(plan: str) -> str:
    """Map DB plan names (FREE/PRO/ENTERPRISE) to PLAN_LIMITS keys."""
    if not plan:
        return "starter"
    p = str(plan).strip().lower()
    if p in ("free", "starter"):
        return "starter"
    if p == "pro":
        return "pro"
    if p == "enterprise":
        return "enterprise"
    return "starter"

# ==================== DATABASE ====================
client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]
meetings_col = db["meetings"]
localities_col = db["localities"]
users_col = db["users"]
tenants_col = db["tenants"]
invitations_col = db["invitations"]


# ==================== AUTH HELPERS ====================
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()


def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed.encode())


def create_token(user_id: str, email: str) -> str:
    from datetime import timedelta
    payload = {
        "sub": user_id,
        "email": email,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS),
        "iat": datetime.now(timezone.utc),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)


security = HTTPBearer()


async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    """Extract and validate JWT from Authorization header (FastAPI Depends pattern)."""
    token = credentials.credentials
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user = await users_col.find_one({"_id": ObjectId(payload["sub"])})
        if not user:
            raise HTTPException(status_code=401, detail="Utilizator inexistent")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expirat")
    except (jwt.InvalidTokenError, Exception):
        raise HTTPException(status_code=401, detail="Token invalid")


async def verify_meeting_ownership(meeting_id: str, user: dict) -> dict:
    """Fetch meeting and verify access (owner, same tenant, or admin). Returns meeting doc."""
    try:
        meeting = await meetings_col.find_one({"_id": ObjectId(meeting_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="ID meeting invalid")
    if not meeting:
        raise HTTPException(status_code=404, detail="Ședința nu a fost găsită")

    # Admin role bypasses ownership checks
    user_role = user.get("role", "member")
    if user_role == "admin":
        return meeting

    user_id = str(user["_id"])
    user_tenant = str(user["tenant_id"]) if user.get("tenant_id") else None
    meeting_user = meeting.get("user_id")
    meeting_tenant = meeting.get("tenant_id")

    # Owner match
    if meeting_user and meeting_user == user_id:
        return meeting

    # Same tenant match (for enterprise multi-user scenarios)
    if user_tenant and meeting_tenant and user_tenant == meeting_tenant:
        return meeting

    raise HTTPException(status_code=403, detail="Nu ai acces la această ședință")


def build_meetings_scope_query(user: dict) -> dict:
    """Build a MongoDB query that scopes meetings to what the user can see.
    - admin: sees all
    - enterprise members (with tenant_id): see tenant-wide meetings
    - regular users: see only their own meetings
    """
    user_role = user.get("role", "member")
    if user_role == "admin":
        return {}

    user_id = str(user["_id"])
    user_tenant = str(user["tenant_id"]) if user.get("tenant_id") else None

    if user_tenant:
        # Enterprise user — see own meetings OR any meeting in same tenant
        return {"$or": [
            {"user_id": user_id},
            {"tenant_id": user_tenant},
        ]}

    # Solo user — own meetings only
    return {"user_id": user_id}


def require_role(*allowed_roles: str):
    """Dependency factory: require user to have one of the given roles (or be admin)."""
    async def checker(user: dict = Depends(get_current_user)) -> dict:
        role = user.get("role", "member")
        if role == "admin" or role in allowed_roles:
            return user
        raise HTTPException(
            status_code=403,
            detail=f"Acțiune restricționată. Necesită unul din rolurile: {', '.join(allowed_roles)}"
        )
    return checker


# ==================== PLAN LIMIT HELPERS ====================
async def reset_monthly_if_needed(user: dict) -> dict:
    """Reset minutes_used_this_month if we're in a new month (defensive — the
    monthly_usage_reset_job task is the canonical reset, this is a safety net
    in case the server was offline at the rollover moment)."""
    last_reset = user.get("last_monthly_reset")
    now = datetime.now(timezone.utc)
    if last_reset is None or (isinstance(last_reset, datetime) and (last_reset.month != now.month or last_reset.year != now.year)):
        await users_col.update_one(
            {"_id": user["_id"]},
            {"$set": {"minutes_used_this_month": 0.0, "last_monthly_reset": now}}
        )
        user["minutes_used_this_month"] = 0.0
        user["last_monthly_reset"] = now
    return user


def get_plan_quota(user: dict) -> tuple[str, float, float]:
    """Return (plan_key, used_minutes, limit_minutes) for a user."""
    plan_key = normalize_plan(user.get("plan"))
    limit = float(PLAN_LIMITS[plan_key]["minutes"])
    used = float(user.get("minutes_used_this_month", 0.0) or 0.0)
    return plan_key, used, limit


# ==================== EMAIL HELPERS ====================
# Verification link points to the backend, which marks email_verified=True
# and then 302-redirects to the deep link `meetingsro://email-verified`.
API_PUBLIC_URL = os.environ.get("API_PUBLIC_URL", "https://meetings-ro-api.onrender.com")
EMAIL_FROM = "Meetings.ro <noreply@meetings-ro.app>"


def build_verification_email_html(verification_url: str, name: str) -> str:
    """Branded HTML for the verification email — navy header, ivory body, orange accent."""
    greeting = f"Bună ziua, {name}!" if name else "Bună ziua!"
    return f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="margin:0;padding:0;background:#FAF8F3;font-family:Georgia,serif;">
  <table width="100%" cellpadding="0" cellspacing="0" style="background:#FAF8F3;padding:40px 20px;">
    <tr><td align="center">
      <table width="560" cellpadding="0" cellspacing="0" style="background:#FFFFFF;border-radius:16px;overflow:hidden;box-shadow:0 2px 8px rgba(0,0,0,0.06);">
        <tr>
          <td style="background:#1B2A4A;padding:32px 40px;text-align:center;">
            <span style="font-size:36px;font-weight:700;color:#FFFFFF;letter-spacing:-1px;">M</span>
            <span style="display:block;font-size:12px;color:rgba(255,255,255,0.6);margin-top:4px;letter-spacing:2px;text-transform:uppercase;">Meetings.ro</span>
          </td>
        </tr>
        <tr>
          <td style="padding:40px 40px 32px;">
            <h1 style="margin:0 0 16px;font-size:24px;color:#1B2A4A;font-weight:600;">{greeting}</h1>
            <p style="margin:0 0 24px;font-size:16px;color:#4A5568;line-height:1.6;">
              Îți mulțumim că te-ai înregistrat pe <strong>Meetings.ro</strong>.
              Pentru a-ți activa contul, te rugăm să verifici adresa de email.
            </p>
            <table cellpadding="0" cellspacing="0" style="margin:0 0 24px;">
              <tr><td style="background:#1B2A4A;border-radius:8px;padding:14px 32px;">
                <a href="{verification_url}"
                   style="color:#FFFFFF;text-decoration:none;font-size:16px;font-weight:600;font-family:Arial,sans-serif;">
                  Verifică adresa de email →
                </a>
              </td></tr>
            </table>
            <p style="margin:0 0 8px;font-size:13px;color:#9CA3AF;">Sau copiază acest link în browser:</p>
            <p style="margin:0 0 24px;font-size:12px;color:#E87040;word-break:break-all;">{verification_url}</p>
            <p style="margin:0;font-size:13px;color:#9CA3AF;line-height:1.5;">
              Link-ul expiră în <strong>24 de ore</strong>.
              Dacă nu ai creat acest cont, poți ignora acest email.
            </p>
          </td>
        </tr>
        <tr>
          <td style="background:#F8F7F4;padding:20px 40px;text-align:center;border-top:1px solid #E5E7EB;">
            <p style="margin:0;font-size:12px;color:#9CA3AF;">© 2026 Meetings.ro · NEDEROV COMEX SRL</p>
          </td>
        </tr>
      </table>
    </td></tr>
  </table>
</body>
</html>"""


async def send_verification_email(email: str, token: str, name: str = "") -> bool:
    """
    Send the email-verification link via Resend.
    Non-blocking — caller should `asyncio.create_task(...)` so registration
    is not delayed by SMTP / network latency.
    """
    if not os.environ.get("RESEND_API_KEY"):
        print(f"[EMAIL] ⚠️ RESEND_API_KEY not set — cannot send verification to {email}")
        return False

    verification_url = f"{API_PUBLIC_URL}/api/auth/verify-email?token={token}"
    try:
        params = {
            "from": EMAIL_FROM,
            "to": [email],
            "subject": "Verifică-ți adresa de email — Meetings.ro",
            "html": build_verification_email_html(verification_url, name),
        }
        # resend SDK is sync — push to a thread so we don't block the event loop
        await asyncio.to_thread(resend.Emails.send, params)
        print(f"[Email] Verification email sent to {email}")
        return True
    except Exception as e:
        print(f"[Email] Failed to send verification email to {email}: {e}")
        return False


def is_email_verified(user: dict) -> bool:
    """Check verification status. Reads either field for backward compat."""
    return bool(user.get("email_verified") or user.get("is_verified"))


# ==================== HELPERS ====================
def serialize_doc(doc):
    """Convert MongoDB document to JSON-serializable dict."""
    if doc is None:
        return None
    result = {}
    for key, value in doc.items():
        if isinstance(value, ObjectId):
            result[key] = str(value)
        elif isinstance(value, datetime):
            result[key] = value.isoformat()
        elif isinstance(value, list):
            result[key] = [serialize_doc(v) if isinstance(v, dict) else (str(v) if isinstance(v, ObjectId) else v) for v in value]
        elif isinstance(value, dict):
            result[key] = serialize_doc(value)
        else:
            result[key] = value
    return result


async def ensure_indexes():
    """Create MongoDB indexes."""
    await meetings_col.create_index("locality")
    await meetings_col.create_index("status")
    await meetings_col.create_index("created_at")
    await meetings_col.create_index([("title", "text"), ("transcript", "text"), ("locality", "text")])


# Default localities (folders)
DEFAULT_LOCALITIES = ["Chilia Veche", "Crișan", "C.A.Rosetti", "Maliuc", "Beștepe"]


async def seed_default_localities():
    """Ensure default locality folders exist."""
    for name in DEFAULT_LOCALITIES:
        await localities_col.update_one(
            {"name": name},
            {"$setOnInsert": {"name": name, "created_at": datetime.now(timezone.utc), "is_default": True}},
            upsert=True
        )
    print(f"[GAL] Default localities seeded: {DEFAULT_LOCALITIES}")


async def monthly_usage_reset_job():
    """Resets minutes_used_this_month for all users on the 1st of each month at 00:00 UTC."""
    while True:
        now = datetime.now(timezone.utc)
        if now.month == 12:
            next_reset = datetime(now.year + 1, 1, 1, tzinfo=timezone.utc)
        else:
            next_reset = datetime(now.year, now.month + 1, 1, tzinfo=timezone.utc)

        wait_seconds = (next_reset - now).total_seconds()
        print(f"[Reset] Next usage reset in {wait_seconds/3600:.1f}h ({next_reset.date()})")
        try:
            await asyncio.sleep(wait_seconds)
        except asyncio.CancelledError:
            print("[Reset] monthly_usage_reset_job cancelled")
            return

        try:
            result = await users_col.update_many(
                {},
                {"$set": {"minutes_used_this_month": 0.0, "last_monthly_reset": datetime.now(timezone.utc)}}
            )
            print(f"[Reset] ✅ Monthly reset — {result.modified_count} users reset")
        except Exception as e:
            print(f"[Reset] ⚠️ Failed monthly reset: {e}")
            # short sleep to avoid hammering on persistent error
            await asyncio.sleep(60)


async def create_demo_account():
    """Create a demo account for Apple Review if it doesn't exist."""
    demo_email = "demo@meetings.ro"
    existing = await users_col.find_one({"email": demo_email})
    if not existing:
        await users_col.insert_one({
            "email": demo_email,
            "password_hash": hash_password("Demo2026!"),
            "name": "Demo User",
            "plan": "PRO",
            "email_verified": True,
            "minutes_used_this_month": 0.0,
            "last_monthly_reset": datetime.now(timezone.utc),
            "created_at": datetime.now(timezone.utc),
        })
        print("[DEMO] Demo account created: demo@meetings.ro / Demo2026!")


@app.on_event("startup")
async def startup():
    await ensure_indexes()
    # User indexes
    await users_col.create_index("email", unique=True)
    await users_col.create_index("reset_token", sparse=True)
    await users_col.create_index("verify_token", sparse=True)
    # Meeting indexes for user-scoped queries
    await meetings_col.create_index("user_id")
    await meetings_col.create_index([("user_id", 1), ("created_at", -1)])
    await meetings_col.create_index("tenant_id", sparse=True)
    await users_col.create_index("tenant_id", sparse=True)
    # Tenants + Invitations indexes
    await tenants_col.create_index("created_by")
    await invitations_col.create_index("token", unique=True)
    await invitations_col.create_index("expires_at", expireAfterSeconds=0)
    await seed_default_localities()
    await create_demo_account()
    # Background: reset minute quotas on the 1st of each month at 00:00 UTC
    asyncio.create_task(monthly_usage_reset_job())
    print("[Meetings.ro] Server started. Indexes ensured.")


# ==================== AUTH ENDPOINTS ====================

class RegisterRequest(BaseModel):
    name: str
    email: str
    password: str
    company: Optional[str] = None
    billing_name: Optional[str] = None     # Full legal name / company name for invoices
    billing_cui: Optional[str] = None      # CUI / VAT code (Romanian companies)
    billing_address: Optional[str] = None  # Billing address for invoices


class LoginRequest(BaseModel):
    email: str
    password: str


@app.post("/api/auth/register")
@limiter.limit("5/minute")
async def auth_register(request: Request, req: RegisterRequest):
    """Register a new user with email verification."""
    # Check if email already exists
    existing = await users_col.find_one({"email": req.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email deja înregistrat")

    if len(req.password) < 6:
        raise HTTPException(status_code=400, detail="Parola trebuie să aibă minim 6 caractere")

    # Generate verification token
    verification_token = secrets.token_urlsafe(32)

    now = datetime.now(timezone.utc)
    user_doc = {
        "name": req.name,
        "email": req.email,
        "password_hash": hash_password(req.password),
        "company": req.company,
        "billing_name": req.billing_name or req.name,
        "billing_cui": req.billing_cui or "",
        "billing_address": req.billing_address or "",
        "plan": "FREE",
        "role": "member",
        "tenant_id": None,
        "email_verified":             False,
        "email_verification_token":   verification_token,
        "email_verification_sent_at": now,
        "minutes_used_this_month": 0.0,
        "last_monthly_reset": now,
        "keep_audio": False,  # GDPR-compliant default — audio is auto-deleted after processing
        "created_at": now,
    }
    result = await users_col.insert_one(user_doc)
    user_id = str(result.inserted_id)

    # Send verification email — non-blocking. Don't fail registration if email
    # delivery is slow or temporarily down.
    asyncio.create_task(send_verification_email(
        email=req.email, token=verification_token, name=req.name,
    ))

    # Auto-login so the mobile app can show the "verify your email" screen
    # immediately. Verification gate is enforced at action time (upload/create).
    token = create_token(user_id, req.email)

    return {
        "token": token,
        "user": {
            "_id": user_id,
            "name": req.name,
            "email": req.email,
            "company": req.company,
            "plan": "FREE",
            "role": "member",
            "tenant_id": None,
            "email_verified": False,
            "minutes_used_this_month":  0.0,
            "minutes_limit_this_month": PLAN_LIMITS["starter"]["minutes"],
            "created_at": now.isoformat(),
        }
    }


@app.post("/api/auth/login")
@limiter.limit("10/minute")
async def auth_login(request: Request, req: LoginRequest):
    """Login with email and password."""
    user = await users_col.find_one({"email": req.email})
    if not user:
        raise HTTPException(status_code=401, detail="Email sau parolă incorectă")

    if not verify_password(req.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Email sau parolă incorectă")

    # NOTE: We deliberately do NOT auto-verify here. The mobile app reads
    # email_verified from the response and routes to /verify-email if False.
    # Verification gate is enforced at action time (upload, create_meeting).

    user_id = str(user["_id"])
    token = create_token(user_id, req.email)

    plan_key, used_minutes, limit_minutes = get_plan_quota(user)
    return {
        "token": token,
        "user": {
            "_id": user_id,
            "name": user.get("name", ""),
            "email": user["email"],
            "company": user.get("company"),
            "plan": user.get("plan", "FREE"),
            "role": user.get("role", "member"),
            "tenant_id": str(user["tenant_id"]) if user.get("tenant_id") else None,
            "email_verified": is_email_verified(user),
            "minutes_used_this_month":  round(used_minutes, 1),
            "minutes_limit_this_month": int(limit_minutes),
            "created_at": user["created_at"].isoformat() if isinstance(user.get("created_at"), datetime) else str(user.get("created_at", "")),
        }
    }


# ---- ADMIN: Require admin role ----
async def require_admin(user: dict = Depends(get_current_user)) -> dict:
    """Dependency that requires the user to have admin role."""
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Acces permis doar administratorilor")
    return user


@app.delete("/api/admin/delete-user")
async def admin_delete_user(email: str = Query(...), admin: dict = Depends(require_admin)):
    """Delete a user by email. Requires admin role."""
    result = await users_col.delete_one({"email": email})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    return {"message": f"User {email} deleted", "deleted": result.deleted_count}


@app.delete("/api/admin/delete-all-users")
async def admin_delete_all_users(admin: dict = Depends(require_admin)):
    """Delete ALL users. Requires admin role."""
    result = await users_col.delete_many({})
    return {"message": f"All users deleted", "deleted": result.deleted_count}


# ⚠️ TEMPORARY — remove after testing.
# Wipes EVERY user and meeting. Behind the same shared-secret guard as
# cleanup-test-users so we don't accidentally nuke production.
@app.delete("/api/admin/cleanup-all-data")
async def cleanup_all_data(secret: str):
    """Temporary endpoint to wipe all test data. DELETE AFTER USE."""
    expected = os.environ.get("ADMIN_CLEANUP_SECRET", "")
    if not expected or secret != expected:
        raise HTTPException(status_code=403, detail="Forbidden")

    users_result    = await users_col.delete_many({})
    meetings_result = await meetings_col.delete_many({})

    print(f"[CleanupAll] Deleted {users_result.deleted_count} users, "
          f"{meetings_result.deleted_count} meetings")

    return {
        "deleted_users":    users_result.deleted_count,
        "deleted_meetings": meetings_result.deleted_count,
        "status":           "clean",
    }


# ⚠️ TEMPORARY — remove after fresh-flow testing is done.
# Auth via shared secret rather than admin role since no admin user exists yet.
@app.delete("/api/admin/cleanup-test-users")
async def cleanup_test_users(secret: str):
    """Temporary endpoint to clean test users + their meetings. DELETE AFTER USE."""
    expected = os.environ.get("ADMIN_CLEANUP_SECRET", "")
    if not expected or secret != expected:
        raise HTTPException(status_code=403, detail="Forbidden")

    test_emails = [
        "vladgrigorov1@gmail.com",
    ]

    # Look up user IDs first so we can purge their meetings too
    users = await users_col.find(
        {"email": {"$in": test_emails}},
        {"_id": 1, "email": 1},
    ).to_list(None)

    # meetings.user_id is stored as STRING (see create_meeting), so stringify
    user_id_strs = [str(u["_id"]) for u in users]

    meetings_result = await meetings_col.delete_many(
        {"user_id": {"$in": user_id_strs}}
    ) if user_id_strs else None

    users_result = await users_col.delete_many(
        {"email": {"$in": test_emails}}
    )

    print(f"[CleanupTestUsers] Deleted {users_result.deleted_count} user(s), "
          f"{meetings_result.deleted_count if meetings_result else 0} meeting(s) "
          f"for emails {test_emails}")

    return {
        "deleted_users":    users_result.deleted_count,
        "deleted_meetings": meetings_result.deleted_count if meetings_result else 0,
        "emails":           test_emails,
    }


@app.get("/api/auth/me")
async def auth_me(user: dict = Depends(get_current_user)):
    """Get current authenticated user."""
    user = await reset_monthly_if_needed(user)
    plan_key, used_minutes, limit_minutes = get_plan_quota(user)
    return {
        "user": {
            "_id": str(user["_id"]),
            "name": user.get("name", ""),
            "email": user["email"],
            "company": user.get("company"),
            "plan": user.get("plan", "FREE"),
            "email_verified":            is_email_verified(user),
            "minutes_used_this_month":   round(used_minutes, 1),
            "minutes_limit_this_month":  int(limit_minutes),
            "created_at": user["created_at"].isoformat() if isinstance(user.get("created_at"), datetime) else str(user.get("created_at", "")),
            "tenant_id": user.get("tenant_id"),   # org membership — drives SlideDrawer menu
            "role": user.get("role"),              # admin / member — drives Admin badge
            "keep_audio": bool(user.get("keep_audio", False)),  # audio retention preference
        }
    }


# ==================== USER SETTINGS ====================

class AudioRetentionUpdate(BaseModel):
    keep_audio: bool


@app.patch("/api/users/audio-retention")
async def update_audio_retention(
    payload: AudioRetentionUpdate,
    current_user: dict = Depends(get_current_user),
):
    """
    Update the user's audio retention preference.
    - keep_audio = False (default): audio is deleted after processing (GDPR-compliant)
    - keep_audio = True: audio is kept on the server until the user deletes the meeting
    """
    await users_col.update_one(
        {"_id": current_user["_id"]},
        {"$set": {"keep_audio": bool(payload.keep_audio)}},
    )
    action = "păstrat" if payload.keep_audio else "șters automat"
    print(f"[Settings] User {current_user.get('email')} — audio va fi {action} după procesare")
    return {"ok": True, "keep_audio": bool(payload.keep_audio)}


# ==================== TOKEN REFRESH ====================

class RefreshRequest(BaseModel):
    token: str


@app.post("/api/auth/refresh")
async def auth_refresh(req: RefreshRequest):
    """Refresh an expiring JWT token. Returns new token if current one is still valid."""
    try:
        payload = jwt.decode(req.token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user = await users_col.find_one({"_id": ObjectId(payload["sub"])})
        if not user:
            raise HTTPException(status_code=401, detail="Utilizator inexistent")
        new_token = create_token(str(user["_id"]), user["email"])
        return {"token": new_token}
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expirat — reconectează-te")
    except (jwt.InvalidTokenError, Exception):
        raise HTTPException(status_code=401, detail="Token invalid")


# ==================== PASSWORD RESET ====================

class ForgotPasswordRequest(BaseModel):
    email: str


class ResetPasswordRequest(BaseModel):
    token: str
    new_password: str


@app.post("/api/auth/forgot-password")
@limiter.limit("3/minute")
async def forgot_password(request: Request, req: ForgotPasswordRequest):
    """Send a password reset email. Always returns success to prevent email enumeration."""
    user = await users_col.find_one({"email": req.email})
    if not user:
        return {"message": "Dacă adresa există, vei primi un email de resetare."}

    reset_token = secrets.token_urlsafe(32)
    expires = datetime.now(timezone.utc) + timedelta(hours=1)

    await users_col.update_one(
        {"_id": user["_id"]},
        {"$set": {"reset_token": reset_token, "reset_token_expires": expires}}
    )

    # Send reset email via Resend
    try:
        reset_html = f"""
        <div style="font-family: 'DM Sans', Arial, sans-serif; max-width: 480px; margin: 0 auto; padding: 40px 20px;">
            <h1 style="color: #1B2A4A; font-size: 24px;">Resetare parolă</h1>
            <p style="color: #444; font-size: 16px;">Ai solicitat resetarea parolei pentru contul tău Meetings.ro.</p>
            <p style="color: #444; font-size: 16px;">Folosește codul de mai jos în aplicație:</p>
            <div style="background: #F3F4F6; padding: 20px; border-radius: 8px; text-align: center; margin: 20px 0;">
                <span style="font-size: 28px; font-weight: bold; color: #1B2A4A; letter-spacing: 2px;">{reset_token[:8]}</span>
            </div>
            <p style="color:#888; font-size:13px;">Codul expiră în 1 oră.</p>
            <hr style="border:none; border-top:1px solid #eee; margin: 30px 0;">
            <p style="color:#888; font-size:12px;">Meetings.ro — Transcriere și sinteză AI pentru orice domeniu</p>
        </div>
        """
        params: resend.Emails.SendParams = {
            "from": "Meetings.ro <onboarding@resend.dev>",
            "to": [req.email],
            "subject": "Resetare parolă Meetings.ro",
            "html": reset_html,
        }
        resend.Emails.send(params)
    except Exception as e:
        print(f"[PASSWORD RESET] Failed to send reset email: {e}")

    return {"message": "Dacă adresa există, vei primi un email de resetare."}


@app.post("/api/auth/reset-password")
@limiter.limit("5/minute")
async def reset_password(request: Request, req: ResetPasswordRequest):
    """Reset password using a valid reset token."""
    if len(req.new_password) < 6:
        raise HTTPException(status_code=400, detail="Parola trebuie să aibă minim 6 caractere")

    # Find user by reset_token (match first 8 chars or full token)
    user = await users_col.find_one({"reset_token": {"$regex": f"^{re.escape(req.token)}"}})
    if not user:
        raise HTTPException(status_code=400, detail="Cod invalid sau expirat")

    # Check expiry
    expires = user.get("reset_token_expires")
    if expires:
        if expires.tzinfo is None:
            expires = expires.replace(tzinfo=timezone.utc)
        if datetime.now(timezone.utc) > expires:
            raise HTTPException(status_code=400, detail="Cod expirat. Solicită un nou email de resetare.")

    # Update password and clear reset token
    await users_col.update_one(
        {"_id": user["_id"]},
        {
            "$set": {"password_hash": hash_password(req.new_password)},
            "$unset": {"reset_token": "", "reset_token_expires": ""}
        }
    )

    return {"message": "Parola a fost resetată cu succes. Te poți autentifica."}


# ==================== EMAIL VERIFICATION ENDPOINTS ====================

# Returns a small HTML page that tries the deep link AND shows a fallback,
# in case the user opens the email on a desktop where meetingsro:// is not registered.
def _verification_redirect_page(status: str) -> str:
    # Triple slash → empty host, /email-verified path. Expo Router routes
    # path-style deep links via its file-based routing.
    deep_link = f"meetingsro:///email-verified?status={status}"
    if status == "success":
        title = "✓ Email verificat"
        body  = "Contul tău Meetings.ro este activ. Te poți întoarce în aplicație."
    elif status == "already_verified":
        title = "Email deja verificat"
        body  = "Contul a fost deja confirmat. Te poți întoarce în aplicație."
    else:
        title = "Link invalid sau expirat"
        body  = "Solicită un email nou din aplicație."
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Meetings.ro</title>
<script>setTimeout(function(){{ window.location.href = "{deep_link}"; }}, 200);</script>
</head><body style="font-family:Georgia,serif;background:#FAF8F3;text-align:center;padding:60px 20px;">
<div style="max-width:480px;margin:0 auto;background:#FFFFFF;border-radius:16px;padding:40px;box-shadow:0 2px 8px rgba(0,0,0,0.06);">
  <h1 style="color:#1B2A4A;margin:0 0 16px;">{title}</h1>
  <p style="color:#4A5568;font-size:16px;line-height:1.6;margin:0 0 24px;">{body}</p>
  <a href="{deep_link}" style="display:inline-block;background:#1B2A4A;color:#FFFFFF;padding:14px 28px;border-radius:8px;text-decoration:none;font-weight:600;">Deschide aplicația</a>
</div>
</body></html>"""


@app.get("/api/auth/verify-email")
async def verify_email(token: str = Query(...)):
    """
    Verifies email using the token from the verification link.
    Marks user as verified and redirects to the mobile app via deep link.
    """
    if not token:
        raise HTTPException(status_code=400, detail="Token lipsă")

    # Look up user by token (regardless of verification status, so we can
    # tell "already verified" apart from "invalid token")
    user = await users_col.find_one({"email_verification_token": token})

    if not user:
        # Maybe they clicked the link twice — token was unset after first use
        return HTMLResponse(_verification_redirect_page("already_verified"), status_code=200)

    if is_email_verified(user):
        return HTMLResponse(_verification_redirect_page("already_verified"), status_code=200)

    # 24h expiry check
    sent_at = user.get("email_verification_sent_at")
    if isinstance(sent_at, datetime):
        if sent_at.tzinfo is None:
            sent_at = sent_at.replace(tzinfo=timezone.utc)
        if datetime.now(timezone.utc) - sent_at > timedelta(hours=24):
            return HTMLResponse(_verification_redirect_page("expired"), status_code=400)

    await users_col.update_one(
        {"_id": user["_id"]},
        {
            "$set": {"email_verified": True, "email_verified_at": datetime.now(timezone.utc)},
            "$unset": {
                "email_verification_token":   "",
                "email_verification_sent_at": "",
                # also clear legacy fields if present
                "verify_token":         "",
                "verify_token_expires": "",
                "is_verified":          "",
            }
        }
    )
    print(f"[Email] Verified: {user['email']}")
    return HTMLResponse(_verification_redirect_page("success"), status_code=200)


# Keep the old /api/auth/verify path working — it just forwards to the new one
@app.get("/api/auth/verify")
async def verify_email_legacy(token: str = Query(...)):
    return await verify_email(token=token)


@app.post("/api/auth/resend-verification")
async def resend_verification(current_user: dict = Depends(get_current_user)):
    """Resends verification email for the currently logged-in user."""
    if is_email_verified(current_user):
        return {"ok": True, "message": "Email deja verificat"}

    # Rate limit: max 3 emails per hour
    now = datetime.now(timezone.utc)
    one_hour_ago = now - timedelta(hours=1)
    sent_times = current_user.get("verification_emails_sent", [])
    recent_sends = [t for t in sent_times if isinstance(t, datetime) and (t.replace(tzinfo=timezone.utc) if t.tzinfo is None else t) > one_hour_ago]
    if len(recent_sends) >= 3:
        raise HTTPException(
            status_code=429,
            detail="Prea multe încercări. Așteaptă o oră înainte de a solicita un alt email."
        )

    new_token = secrets.token_urlsafe(32)
    await users_col.update_one(
        {"_id": current_user["_id"]},
        {
            "$set": {
                "email_verification_token":   new_token,
                "email_verification_sent_at": now,
            },
            "$push": {"verification_emails_sent": now},
        }
    )

    asyncio.create_task(send_verification_email(
        email=current_user["email"], token=new_token, name=current_user.get("name", ""),
    ))
    return {"ok": True, "message": "Email de verificare retrimis"}

    return {"message": "Email de confirmare retrimis. Verifică inbox-ul."}


# ==================== DELETE ACCOUNT ====================

@app.delete("/api/auth/delete-account")
async def delete_account(current_user: dict = Depends(get_current_user)):
    """Delete user account and all associated data. Required by Apple App Store."""
    import shutil
    user_id = str(current_user["_id"])

    # Delete all user's meetings
    await meetings_col.delete_many({"user_id": user_id})

    # Delete user record
    await users_col.delete_one({"_id": current_user["_id"]})

    # Delete uploaded audio files
    user_uploads_path = UPLOAD_DIR / user_id
    if user_uploads_path.exists():
        shutil.rmtree(str(user_uploads_path), ignore_errors=True)

    return {"message": "Contul și toate datele asociate au fost șterse definitiv."}


# ==================== PUSH NOTIFICATIONS ====================

class PushTokenRequest(BaseModel):
    token: str

@app.post("/api/users/push-token")
async def save_push_token(req: PushTokenRequest, current_user: dict = Depends(get_current_user)):
    """Save Expo push token for the current user."""
    await users_col.update_one(
        {"_id": current_user["_id"]},
        {"$set": {"push_token": req.token}}
    )
    return {"ok": True}


async def notify_user_meeting_done(user_id: str, meeting_title: str):
    """Send push notification when meeting processing is complete."""
    try:
        user = await users_col.find_one({"_id": ObjectId(user_id)})
        if not user:
            return
        push_token = user.get("push_token")
        if not push_token:
            return
        import httpx
        async with httpx.AsyncClient() as client:
            await client.post("https://exp.host/--/api/v2/push/send", json={
                "to": push_token,
                "title": "Raport gata",
                "body": f"{meeting_title or 'Întâlnirea ta'} a fost procesată.",
                "sound": "default",
            })
    except Exception as e:
        print(f"[Push] Failed to send notification: {e}")


# ==================== USAGE ENDPOINT ====================
@app.get("/api/users/me/usage")
async def get_user_usage(user: dict = Depends(get_current_user)):
    """Get current user's plan usage stats (minute-based)."""
    user = await reset_monthly_if_needed(user)
    plan_key, used_minutes, limit_minutes = get_plan_quota(user)
    remaining = max(0.0, limit_minutes - used_minutes)
    percentage = round((used_minutes / limit_minutes) * 100, 1) if limit_minutes > 0 else 0.0
    return {
        "plan": user.get("plan", "FREE"),
        "plan_key": plan_key,
        "minutes_used":      round(used_minutes, 1),
        "minutes_limit":     int(limit_minutes),
        "minutes_remaining": round(remaining, 1),
        "percentage":        min(100.0, percentage),
        "label":             PLAN_LIMITS[plan_key]["label"],
    }


# ==================== MODELS ====================
class ActionItemModel(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    text: str
    owner: Optional[str] = None
    deadline: Optional[str] = None
    completed: bool = False


class MeetingCreate(BaseModel):
    title: Optional[str] = None
    locality: Optional[str] = None
    vertical_type: Optional[str] = "GAL"


class MeetingUpdate(BaseModel):
    title: Optional[str] = None
    locality: Optional[str] = None


# ==================== DIARIZATION CONSTANTS ====================
MAX_DIARIZATION_CHUNK = 8000   # chars per Claude call
CHUNK_OVERLAP = 500            # overlap between chunks for context continuity

# TASK 1 — Engineered institutional diarization prompt
DIARIZATION_SYSTEM_PROMPT = """Ești un expert în analiza transcrierilor din ședințe oficiale românești.
Primești o transcriere cu timestamps dintr-o ședință instituțională (consiliu local, tribunal, spital, universitate, primărie).

SARCINA: Identifică cu precizie maximă fiecare schimbare de vorbitor și atribuie corect fiecare replică.

═══ REGULI DE IDENTIFICARE ════════════════════════════════════════

1. PREZENTĂRI EXPLICITE (prioritate maximă):
   - Dacă cineva spune "Eu sunt [Nume]", "Mă numesc [Nume]", "Din partea [instituției]" → folosește acel nume de acum încolo
   - Dacă menționează funcția: "în calitate de primar", "ca secretar general" → adaugă funcția în câmpul "role"
   - Dacă președintele ședinței se identifică la deschidere → păstrează-l ca speaker consistent

2. FRAZE DE TRANZIȚIE (indică schimbare de vorbitor):
   - "Vă mulțumesc", "Mulțumesc pentru cuvânt" → vorbitor curent cedează cuvântul
   - "Dacă îmi permiteți", "Solicit cuvântul", "Cer cuvântul" → vorbitor nou intră
   - "Domnul/Doamna [nume/funcție]" spus de altcineva → urmează intervenția acelei persoane
   - "Supun la vot", "Cine este pentru?", "Împotrivă?", "Abțineri?" → ÎNTOTDEAUNA președinte/primar
   - "Da / Nu / Abținere" ca răspuns izolat → consilier/participant care votează (nu schimbă speaker principal)

3. ROLURI INSTITUȚIONALE (deduse din context):
   - Cel care deschide ședința, citește ordinea de zi, conduce votul → Președinte ședință / Primar
   - Cel care citește un referat sau raport → Secretar / Director / Referent
   - Cei care intervin scurt cu "sunt de acord", "propun", "amendament" → Consilieri
   - Cel care consemnează ("Am înregistrat", "Se consemnează") → Secretar general

4. CONSISTENȚĂ (regula de aur):
   - Dacă ai identificat că Vorbitor 1 e primarul → ORICE frază de vot sau deschidere ulterioară = același Vorbitor 1
   - Nu schimba speaker dacă e aceeași persoană care continuă după o pauză scurtă (<3 secunde)
   - Dacă două segmente consecutive au același pattern de vorbire și nu există semnal de schimbare → același vorbitor

5. CÂND NU ȘTII:
   - Folosește Vorbitor N (număr incremental)
   - Nu inventa nume sau funcții
   - Nu modifica niciun cuvânt din transcriere

═══ STRUCTURI SPECIFICE DE ȘEDINȚĂ ════════════════════════════════

Ședință consiliu local: Primar (conduce) → Secretar (citește) → Consilieri (intervin) → Primar (vot)
Ședință tribunal: Judecător (conduce) → Grefier (consemnează) → Avocați → Părți
Ședință medicală: Director (conduce) → Medici șefi (raportează) → Personal (intervenții)

═══ FORMAT RĂSPUNS ═══════════════════════════════════════════════

Răspunde STRICT cu JSON valid — array de obiecte, fără text în afara JSON-ului:
[
  {
    "speaker": "Primar Ion Popescu",
    "role": "Primar",
    "timestamp": "00:00",
    "text": "Declar deschisă ședința ordinară a Consiliului Local..."
  }
]

REGULI FORMAT:
- "speaker": nume real dacă identificat, altfel "Vorbitor 1", "Vorbitor 2" etc.
- "role": funcția dacă identificată (ex: "Primar", "Secretar general", "Consilier"), altfel null
- "timestamp": MM:SS al primului segment din grupul de replici
- "text": textul exact din transcriere, grupat dacă același vorbitor continuă
- Grupează segmentele consecutive ale aceluiași vorbitor într-un singur obiect
- NU adăuga explicații, NU modifica textul transcrierii
"""


# ── TASK 3: Vertical-aware Whisper prompt ─────────────────────────────
def get_whisper_prompt(vertical_type: str) -> str:
    """Returns the optimal Whisper initial_prompt for the given vertical.
    Injects domain vocabulary to reduce WER for specialist terminology.
    """
    from verticals import get_vertical_config
    try:
        config = get_vertical_config(vertical_type)
        if config.whisper_prompt:
            return config.whisper_prompt
    except Exception:
        pass
    # Universal fallback
    return (
        "Înregistrare audio în limba română. "
        "Vorbire naturală sau formală. "
        "Transcriere exactă, fidelă."
    )


# ── TASK 4: Vertical-aware diarization prompt builder ─────────────────
def build_diarization_prompt(vertical_type: str) -> str:
    """
    Builds a context-aware diarization system prompt for the given vertical.
    Combines the universal institutional rules with vertical-specific hints.
    """
    from verticals import get_vertical_config

    # Universal rules — apply to all verticals
    universal_rules = DIARIZATION_SYSTEM_PROMPT

    # Vertical-specific context appendix
    vertical_context = ""
    try:
        config = get_vertical_config(vertical_type)
        if config.diarization_context:
            vertical_context += f"\nCONTEXT SPECIFIC ACESTUI TIP DE ÎNREGISTRARE:\n{config.diarization_context}\n"
        if config.expected_speakers:
            mn, mx = config.expected_speakers
            if mn == mx == 1:
                vertical_context += (
                    "\nATENȚIE SPECIALĂ: Aceasta este o înregistrare cu un singur vorbitor. "
                    "NU fragmenta în mai mulți vorbitori. Un singur obiect în array.\n"
                )
            else:
                vertical_context += (
                    f"\nNumărul așteptat de vorbitori pentru acest tip de înregistrare: "
                    f"între {mn} și {mx}.\n"
                )
    except Exception:
        pass

    return universal_rules + vertical_context


# ==================== AUDIO PREPROCESSING ====================
async def preprocess_audio(input_path: str) -> str:
    """
    TASK 5a — Normalize audio to 16kHz mono WAV using ffmpeg for optimal Whisper accuracy.
    Applies EBU R128 loudness normalization and resamples to 16kHz mono.
    Returns path to preprocessed WAV, or original path if ffmpeg fails/unavailable.
    """
    # Defensive fallback — if ffmpeg binary is missing, skip preprocessing entirely
    # so the pipeline can still attempt direct transcription with the original file.
    if shutil.which("ffmpeg") is None:
        print("[ffmpeg] ⚠️ not available — skipping preprocess step")
        return input_path

    base = input_path.rsplit('.', 1)[0]
    output_path = f"{base}_preprocessed.wav"
    try:
        proc = await asyncio.create_subprocess_exec(
            'ffmpeg', '-y',
            '-i', input_path,
            '-ar', '16000',        # resample to 16kHz (Whisper native sample rate)
            '-ac', '1',            # mono channel
            '-af', 'loudnorm',     # EBU R128 loudness normalization
            '-f', 'wav',
            output_path,
            stdout=asyncio.subprocess.DEVNULL,
            stderr=asyncio.subprocess.DEVNULL,
        )
        await proc.wait()
        if proc.returncode == 0 and os.path.exists(output_path):
            print(f"[Preprocess] Audio normalized → {output_path}")
            return output_path
        else:
            print(f"[Preprocess] ffmpeg returned {proc.returncode}, using original file")
            return input_path
    except Exception as e:
        print(f"[Preprocess] ffmpeg error: {e} — using original file")
        return input_path


# ==================== AUDIO SPLITTING ====================
CHUNK_DURATION_SECONDS = 600          # 10 minutes per chunk
GROQ_CONCURRENCY_LIMIT = 5            # max parallel Groq API calls


async def split_audio_into_chunks(audio_path: str) -> list:
    """
    Splits any audio file into fixed 10-minute WAV chunks. Always splits —
    no size check. Short files produce 1 chunk; long files produce many.
    Returns list of (chunk_path, time_offset_seconds) tuples.
    All chunk files are tempfile-based — caller must clean them up.
    """
    # Defensive fallback — if ffmpeg/ffprobe are missing, return the original
    # file as a single "chunk" with offset 0 so the pipeline can still attempt
    # a direct transcription instead of failing the whole job.
    if shutil.which("ffmpeg") is None or shutil.which("ffprobe") is None:
        print("[ffmpeg] ⚠️ not available — skipping split, returning original file as single chunk")
        return [(audio_path, 0.0)]

    # Step 1: Get total duration via ffprobe (JSON output)
    total_duration: float = 0.0
    try:
        proc = await asyncio.create_subprocess_exec(
            'ffprobe', '-v', 'quiet',
            '-print_format', 'json',
            '-show_format',
            audio_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.DEVNULL,
        )
        stdout, _ = await proc.communicate()
        duration_data = json.loads(stdout.decode())
        total_duration = float(duration_data["format"]["duration"])
        print(f"[Split] Audio duration: {total_duration:.1f}s ({total_duration / 60:.1f} min)")
    except Exception as e:
        print(f"[Split] ffprobe failed ({e}) — estimating from file size")
        file_size = os.path.getsize(audio_path)
        # Conservative estimate: ~1MB/min for 128kbps m4a
        total_duration = (file_size / (1024 * 1024)) * 60
        print(f"[Split] Estimated duration: {total_duration:.1f}s from {file_size / (1024 * 1024):.1f}MB")

    num_chunks = max(1, math.ceil(total_duration / CHUNK_DURATION_SECONDS))
    print(f"[Split] Splitting into {num_chunks} chunk(s) of {CHUNK_DURATION_SECONDS}s")

    # Step 2: Generate all chunks via ffmpeg
    chunk_paths = []
    for i in range(num_chunks):
        start_time = i * CHUNK_DURATION_SECONDS
        chunk_path = tempfile.mktemp(suffix=f"_chunk_{i:03d}.wav")

        proc = await asyncio.create_subprocess_exec(
            'ffmpeg', '-y',
            '-i', audio_path,
            '-ss', str(start_time),
            '-t', str(CHUNK_DURATION_SECONDS),
            '-ar', '16000',         # Whisper native sample rate
            '-ac', '1',             # mono
            '-c:a', 'pcm_s16le',    # 16-bit PCM WAV
            chunk_path,
            stdout=asyncio.subprocess.DEVNULL,
            stderr=asyncio.subprocess.PIPE,
        )
        _, stderr = await proc.communicate()

        if proc.returncode == 0 and os.path.exists(chunk_path):
            chunk_mb = os.path.getsize(chunk_path) / (1024 * 1024)
            print(f"[Split] Chunk {i + 1}/{num_chunks}: {chunk_mb:.1f}MB (offset: {start_time}s)")
            chunk_paths.append((chunk_path, float(start_time)))
        else:
            err_preview = stderr.decode()[:200] if stderr else "unknown"
            print(f"[Split] ⚠️ Chunk {i + 1}/{num_chunks} failed: {err_preview}")
            # Continue — partial transcription > complete failure

    if not chunk_paths:
        raise RuntimeError(f"Audio splitting failed for all {num_chunks} chunks")

    print(f"[Split] ✅ {len(chunk_paths)}/{num_chunks} chunks ready")
    return chunk_paths


# ==================== AUDIO DURATION ====================
async def get_audio_duration_seconds(file_path: str) -> float:
    """
    Returns audio duration in seconds using ffprobe.
    Returns 0.0 if ffprobe is unavailable or file is unreadable.
    Used for minute-based quota accounting on the user's plan.
    """
    if not shutil.which("ffprobe"):
        return 0.0
    try:
        proc = await asyncio.create_subprocess_exec(
            "ffprobe", "-v", "quiet",
            "-print_format", "json",
            "-show_format",
            file_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        try:
            stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=30)
        except asyncio.TimeoutError:
            try: proc.kill()
            except Exception: pass
            return 0.0
        if proc.returncode != 0 or not stdout:
            return 0.0
        data = json.loads(stdout.decode("utf-8", errors="ignore"))
        return float(data["format"]["duration"])
    except Exception as e:
        print(f"[Duration] ffprobe error: {e}")
        return 0.0


# ==================== AI CLIENTS ====================
# Groq for transcription (Whisper), OpenAI as fallback
groq_client = AsyncOpenAI(
    api_key=GROQ_API_KEY,
    base_url="https://api.groq.com/openai/v1",
) if GROQ_API_KEY else None
openai_client = AsyncOpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
anthropic_client = AsyncAnthropic(api_key=ANTHROPIC_API_KEY)


# ==================== AI PROCESSING ====================
async def transcribe_audio(file_path: str, vertical_type: str = "GENERAL") -> dict:
    """
    Transcribe a single audio chunk via Groq Whisper (primary) or OpenAI (fallback).
    Retries on rate limit (429) and transient server errors with exponential backoff.
    Falls back to OpenAI after all Groq retries are exhausted.
    """
    if not groq_client and not openai_client:
        raise RuntimeError("Nicio cheie API pentru transcriere configurată (GROQ_API_KEY sau OPENAI_API_KEY)")

    prompt = get_whisper_prompt(vertical_type)
    max_retries = 4

    for attempt in range(max_retries):
        try:
            client = groq_client or openai_client
            model = "whisper-large-v3" if groq_client else "whisper-1"

            with open(file_path, "rb") as f:
                try:
                    response = await client.audio.transcriptions.create(
                        model=model,
                        file=f,
                        response_format="verbose_json",
                        language="ro",
                        prompt=prompt,
                        temperature=0.0,
                        timestamp_granularities=["word", "segment"],
                    )
                except Exception as ts_err:
                    if "timestamp_granularities" in str(ts_err):
                        f.seek(0)
                        response = await client.audio.transcriptions.create(
                            model=model,
                            file=f,
                            response_format="verbose_json",
                            language="ro",
                            prompt=prompt,
                            temperature=0.0,
                        )
                    else:
                        raise

            # Parse response
            transcript_text = response.text or ""
            segments = []
            words = []

            if hasattr(response, "segments") and response.segments:
                for seg in response.segments:
                    segments.append(seg if isinstance(seg, dict) else {
                        "start": getattr(seg, "start", 0),
                        "end":   getattr(seg, "end", 0),
                        "text":  getattr(seg, "text", "").strip(),
                    })

            if hasattr(response, "words") and response.words:
                for w in response.words:
                    words.append(w if isinstance(w, dict) else {
                        "word":  getattr(w, "word", ""),
                        "start": getattr(w, "start", 0),
                        "end":   getattr(w, "end", 0),
                    })
                print(f"[Transcribe] {len(words)} word-level timestamps")

            return {"text": transcript_text, "segments": segments, "words": words}

        except Exception as e:
            err_str = str(e)
            is_rate_limit = "429" in err_str or "rate_limit" in err_str.lower()
            is_server_err = any(code in err_str for code in ["500", "502", "503", "529"])

            if attempt < max_retries - 1 and (is_rate_limit or is_server_err):
                wait = (2 ** attempt) + (1 if is_rate_limit else 0)
                label = "rate limit" if is_rate_limit else "server error"
                print(f"[Transcribe] Attempt {attempt + 1} failed ({label}), retrying in {wait}s...")
                await asyncio.sleep(wait)
                continue

            # Last attempt — try OpenAI fallback if Groq was primary
            if groq_client and openai_client and attempt == max_retries - 1:
                print(f"[Transcribe] Groq exhausted after {max_retries} attempts — trying OpenAI fallback")
                try:
                    with open(file_path, "rb") as f:
                        fb_response = await openai_client.audio.transcriptions.create(
                            model="whisper-1",
                            file=f,
                            response_format="verbose_json",
                            language="ro",
                            prompt=prompt,
                            temperature=0.0,
                        )
                    return {"text": fb_response.text or "", "segments": [], "words": []}
                except Exception as fb_err:
                    print(f"[Transcribe] OpenAI fallback failed: {fb_err}")

            raise


async def transcribe_chunk(
    chunk_path: str,
    time_offset: float,
    vertical_type: str,
    semaphore: asyncio.Semaphore,
    chunk_index: int,
    total_chunks: int,
) -> dict:
    """
    Transcribes one chunk under semaphore-controlled concurrency.
    Applies time_offset to all segment/word timestamps.
    Never raises — returns {success: False} on any error so gather() keeps running.
    """
    async with semaphore:
        print(f"[Transcribe] Starting chunk {chunk_index + 1}/{total_chunks} (offset: {time_offset:.0f}s)")
        try:
            result = await transcribe_audio(chunk_path, vertical_type=vertical_type)

            offset_segments = [
                {**seg,
                 "start": seg.get("start", 0) + time_offset,
                 "end":   seg.get("end",   0) + time_offset}
                for seg in result.get("segments", [])
            ]
            offset_words = [
                {**w,
                 "start": w.get("start", 0) + time_offset,
                 "end":   w.get("end",   0) + time_offset}
                for w in result.get("words", [])
            ]

            text = result.get("text", "").strip()
            print(f"[Transcribe] ✅ Chunk {chunk_index + 1}/{total_chunks}: {len(text)} chars")
            return {
                "text":     text,
                "segments": offset_segments,
                "words":    offset_words,
                "offset":   time_offset,
                "success":  True,
            }

        except Exception as e:
            print(f"[Transcribe] ⚠️ Chunk {chunk_index + 1}/{total_chunks} failed: {e}")
            return {
                "text":     "",
                "segments": [],
                "words":    [],
                "offset":   time_offset,
                "success":  False,
            }


async def transcribe_with_chunking(audio_path: str, vertical_type: str = "GENERAL") -> dict:
    """
    Main transcription orchestrator — handles any file size, any duration.
    Pipeline: preprocess (16kHz mono) → always-split into 10-min chunks →
              parallel transcription (≤5 concurrent Groq calls) →
              merge with correct timestamps → cleanup all temp files.
    Returns {text, segments, words, chunks_total, chunks_succeeded}.
    """
    temp_files: list = []

    try:
        # Step 1: Normalize audio (16kHz mono WAV, loudnorm)
        preprocessed_path = await preprocess_audio(audio_path)
        if preprocessed_path != audio_path:
            temp_files.append(preprocessed_path)

        # Step 2: Always split into fixed chunks
        chunks = await split_audio_into_chunks(preprocessed_path)
        for chunk_path, _ in chunks:
            temp_files.append(chunk_path)

        total_chunks = len(chunks)
        print(f"[Pipeline] Transcribing {total_chunks} chunk(s) in parallel (max {GROQ_CONCURRENCY_LIMIT} concurrent)")

        # Step 3: Parallel transcription with semaphore rate-limit guard
        semaphore = asyncio.Semaphore(GROQ_CONCURRENCY_LIMIT)
        tasks = [
            transcribe_chunk(
                chunk_path=chunk_path,
                time_offset=time_offset,
                vertical_type=vertical_type,
                semaphore=semaphore,
                chunk_index=i,
                total_chunks=total_chunks,
            )
            for i, (chunk_path, time_offset) in enumerate(chunks)
        ]
        chunk_results = await asyncio.gather(*tasks)

        # Step 4: Sort by offset (gather doesn't guarantee ordering)
        chunk_results = sorted(chunk_results, key=lambda r: r["offset"])

        # Step 5: Merge
        successful = [r for r in chunk_results if r["success"]]
        failed_count = total_chunks - len(successful)

        if not successful:
            raise RuntimeError("All audio chunks failed transcription — check Groq API key and limits")

        if failed_count > 0:
            print(f"[Pipeline] ⚠️ {failed_count}/{total_chunks} chunks failed — partial transcription")

        merged_text = " ".join(r["text"] for r in successful if r["text"])
        merged_segments = sorted(
            [seg for r in successful for seg in r["segments"]],
            key=lambda s: s.get("start", 0),
        )
        merged_words = sorted(
            [w for r in successful for w in r["words"]],
            key=lambda w: w.get("start", 0),
        )

        print(
            f"[Pipeline] ✅ Transcription complete: {len(merged_text)} chars, "
            f"{len(merged_segments)} segments, {len(merged_words)} words — "
            f"{len(successful)}/{total_chunks} chunks succeeded"
        )
        return {
            "text":               merged_text,
            "segments":           merged_segments,
            "words":              merged_words,
            "chunks_total":       total_chunks,
            "chunks_succeeded":   len(successful),
        }

    finally:
        # Clean up all temp files regardless of success/failure
        for path in temp_files:
            try:
                if path and os.path.exists(path):
                    os.unlink(path)
            except Exception:
                pass


# ==================== SPEAKER DIARIZATION ====================

async def extract_speaker_context(segments_text: str) -> str:
    """
    TASK 2 — First pass: cheap context extraction using claude-haiku-4-5.
    Identifies meeting type, named speakers, and speaker count estimate
    from the first 3000 chars. Returns JSON string injected into main pass.
    """
    try:
        response = await anthropic_client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=500,
            system=(
                "Ești un analizor de transcrieri. Extrage DOAR informațiile explicit menționate.\n"
                "Răspunde în JSON cu structura:\n"
                '{"meeting_type": "...", "speakers_identified": [{"name": "...", "role": "..."}], '
                '"total_speakers_estimate": N}'
            ),
            messages=[{
                "role": "user",
                "content": (
                    "Din această transcriere, extrage:\n"
                    "1. Tipul ședinței (consiliu local / tribunal / medical / altul)\n"
                    "2. Numele și funcțiile persoanelor menționate explicit\n"
                    "3. Estimarea numărului de vorbitori diferiți\n\n"
                    f"Transcriere (primele 3000 caractere):\n{segments_text[:3000]}"
                )
            }]
        )
        return response.content[0].text.strip()
    except Exception as e:
        print(f"[Diarize/Context] Extraction failed: {e}")
        return "{}"


def validate_and_fix_diarization(result: list, segments: list) -> list:
    """
    TASK 3 — Post-process Claude diarization output:
    - Remove empty / too-short entries
    - Normalize "Vorbitor 0" → "Vorbitor 1"
    - Validate MM:SS timestamp format
    - Fallback to single-speaker block if result is empty but segments exist
    """
    if not isinstance(result, list):
        return []

    cleaned = []
    speaker_map: dict = {}

    for item in result:
        if not isinstance(item, dict):
            continue

        text = item.get("text", "").strip()
        if not text or len(text) < 2:
            continue

        speaker = item.get("speaker", "").strip()
        if not speaker:
            speaker = f"Vorbitor {len(speaker_map) + 1}"

        # Normalize "Vorbitor 0" → "Vorbitor 1"
        if re.match(r'^Vorbitor\s+0$', speaker, re.IGNORECASE):
            speaker = "Vorbitor 1"

        if speaker not in speaker_map:
            speaker_map[speaker] = speaker

        # Validate / coerce timestamp to MM:SS
        ts = str(item.get("timestamp", "00:00"))
        if not re.match(r'^\d{2}:\d{2}$', ts):
            ts = "00:00"

        cleaned.append({
            "speaker": speaker_map[speaker],
            "role":    item.get("role") or None,
            "timestamp": ts,
            "text":    text,
        })

    # Sanity fallback: if Claude returned nothing but segments exist
    if not cleaned and segments:
        fallback_text = " ".join(s.get("text", "") for s in segments[:50]).strip()
        if fallback_text:
            return [{
                "speaker":   "Vorbitor 1",
                "role":      None,
                "timestamp": "00:00",
                "text":      fallback_text,
            }]

    return cleaned


async def _diarize_with_claude_nlp(
    segments: list,
    segments_text: str = None,
    vertical_type: str = "GENERAL",
) -> list:
    """
    TASK 1+2+3+4 — Single-call diarization path.
    Two-pass: cheap Haiku context extraction → vertical-aware Sonnet diarization.
    Dynamic token budget based on segment count (min 4000, max 8000).
    """
    if not segments or not anthropic_client:
        return []

    # Build segments_text if caller didn't pre-build it
    if segments_text is None:
        segments_text = ""
        for seg in segments:
            start = seg.get("start", 0)
            text  = seg.get("text", "").strip()
            if text:
                minutes = int(start // 60)
                seconds = int(start % 60)
                segments_text += f"[{minutes:02d}:{seconds:02d}] {text}\n"

    if not segments_text.strip():
        return []

    # PASS 1 — cheap context extraction with Haiku (TASK 2)
    speaker_context = await extract_speaker_context(segments_text)

    # Build enriched user message
    user_message = (
        f"CONTEXT EXTRAS DIN TRANSCRIERE:\n{speaker_context}\n\n"
        f"TRANSCRIERE COMPLETĂ DE DIARIZAT:\n\n{segments_text}"
    )

    # Dynamic token budget (TASK 3): ~60 tokens per segment turn
    estimated_turns = len([s for s in segments if s.get("text", "").strip()])
    max_tokens_needed = min(8000, max(4000, estimated_turns * 60))

    # TASK 4 — vertical-aware system prompt
    system_prompt = build_diarization_prompt(vertical_type)

    try:
        response = await anthropic_client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=max_tokens_needed,
            system=system_prompt,
            messages=[{"role": "user", "content": user_message}]
        )

        response_text = response.content[0].text.strip()
        try:
            result = json.loads(response_text)
        except json.JSONDecodeError:
            json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group(0))
            else:
                print("[Diarize] Could not parse JSON from Claude response")
                return validate_and_fix_diarization([], segments)

        return validate_and_fix_diarization(
            result if isinstance(result, list) else [], segments
        )

    except Exception as e:
        print(f"[Diarize] Error: {e}")
        return []


async def diarize_long_transcript(segments: list, vertical_type: str = "GENERAL") -> list:
    """
    TASK 4 — Chunked diarization for long transcripts (>MAX_DIARIZATION_CHUNK chars).
    Splits into overlapping chunks, diarizes each with speaker-continuity context,
    then merges by deduplicating on timestamp comparison.
    Passes vertical_type to each chunk call for consistent prompt context.
    """
    # Build full formatted text once
    full_text = ""
    for seg in segments:
        start = seg.get("start", 0)
        text  = seg.get("text", "").strip()
        if text:
            minutes = int(start // 60)
            seconds = int(start % 60)
            full_text += f"[{minutes:02d}:{seconds:02d}] {text}\n"

    if not full_text.strip():
        return []

    # Split into chunks at newline boundaries
    chunks = []
    start_idx = 0
    while start_idx < len(full_text):
        end_idx = start_idx + MAX_DIARIZATION_CHUNK
        if end_idx < len(full_text):
            # Snap to last newline to avoid cutting mid-segment
            newline_idx = full_text.rfind('\n', start_idx, end_idx)
            if newline_idx > start_idx:
                end_idx = newline_idx + 1
        chunks.append(full_text[start_idx:end_idx])
        start_idx = end_idx - CHUNK_OVERLAP

    print(f"[Diarize] Long transcript split into {len(chunks)} chunks")

    all_turns: list = []
    previous_speakers: list = []

    for i, chunk in enumerate(chunks):
        # Carry speaker labels from previous chunk for consistency
        context_note = ""
        if previous_speakers:
            speaker_list = ", ".join(set(previous_speakers[-10:]))
            context_note = (
                f"NOTĂ: Aceasta este continuarea ședinței. "
                f"Vorbitorii identificați până acum: {speaker_list}. "
                f"Păstrează aceleași etichete pentru aceiași vorbitori.\n\n"
            )

        chunk_with_context = context_note + chunk

        try:
            response = await anthropic_client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=6000,
                system=build_diarization_prompt(vertical_type),
                messages=[{
                    "role": "user",
                    "content": f"Identifică vorbitorii:\n\n{chunk_with_context}"
                }]
            )
            response_text = response.content[0].text.strip()
            try:
                chunk_result = json.loads(response_text)
            except json.JSONDecodeError:
                json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
                chunk_result = json.loads(json_match.group(0)) if json_match else []

            chunk_result = validate_and_fix_diarization(
                chunk_result if isinstance(chunk_result, list) else [], []
            )

            # Deduplicate overlap: skip turns whose timestamp ≤ last saved turn
            if all_turns and chunk_result:
                last_ts = all_turns[-1].get("timestamp", "00:00")
                chunk_result = [t for t in chunk_result if t.get("timestamp", "99:99") > last_ts]

            all_turns.extend(chunk_result)
            previous_speakers.extend([t["speaker"] for t in chunk_result])
            print(f"[Diarize] Chunk {i + 1}/{len(chunks)} → {len(chunk_result)} turns")

        except Exception as e:
            print(f"[Diarize] Chunk {i + 1}/{len(chunks)} failed: {e}")
            continue

    return all_turns


async def diarize_transcript(
    segments: list,
    words: list = None,
    audio_path: str = None,
    vertical_type: str = "GENERAL",
) -> list:
    """
    Public entry point for speaker diarization.
    Routes to chunked path for long transcripts, single-call for short ones.
    Passes vertical_type through the entire chain for context-aware prompts.
    """
    if not segments or not anthropic_client:
        return []

    # Pre-build segments_text to measure total length
    segments_text = ""
    for seg in segments:
        start = seg.get("start", 0)
        text  = seg.get("text", "").strip()
        if text:
            minutes = int(start // 60)
            seconds = int(start % 60)
            segments_text += f"[{minutes:02d}:{seconds:02d}] {text}\n"

    if not segments_text.strip():
        return []

    # TASK 4 — Route to chunked diarization for long transcripts
    if len(segments_text) > MAX_DIARIZATION_CHUNK:
        print(f"[Diarize] Long transcript ({len(segments_text)} chars) → chunked mode")
        return await diarize_long_transcript(segments, vertical_type=vertical_type)

    return await _diarize_with_claude_nlp(segments, segments_text, vertical_type=vertical_type)


# NOTE: personal_legal is intentionally excluded from auto-detection.
# Users must select it manually. This vertical produces verbatim transcripts
# intended for legal use — auto-assignment without user consent is a liability risk.
async def detect_vertical(transcript_sample: str) -> str:
    """
    TASK 6 — Auto-detect the most appropriate vertical from transcript text.
    Uses claude-haiku-4-5 (cheap, max_tokens=50) for speed and cost efficiency.
    Called only when the user left vertical_type as "GENERAL".
    Returns one of the registered vertical_type keys, or "GENERAL" on failure.
    PERSONAL_LEGAL is excluded — it must be selected explicitly by the user.
    """
    if not transcript_sample or not transcript_sample.strip():
        return "GENERAL"

    VALID_VERTICALS = {
        "GAL", "LEGAL", "HEALTHCARE", "BANKING",
        "JOURNALISM", "STARTUPS", "SOLO_READING", "GENERAL",
    }

    try:
        response = await anthropic_client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=50,
            system=(
                "Clasifică textul într-una din categorii. "
                "Răspunde DOAR cu una din (exact, fără alte cuvinte): "
                "GAL, LEGAL, HEALTHCARE, BANKING, JOURNALISM, STARTUPS, "
                "SOLO_READING, GENERAL."
            ),
            messages=[{
                "role": "user",
                "content": f"Clasifică această transcriere:\n\n{transcript_sample[:1500]}"
            }]
        )
        detected = response.content[0].text.strip().upper()
        return detected if detected in VALID_VERTICALS else "GENERAL"
    except Exception as e:
        print(f"[AutoDetect] Failed: {e}")
        return "GENERAL"


def _parse_claude_json(response_text: str) -> dict:
    """Tolerant JSON extraction from Claude's response (handles bare JSON / fenced blocks)."""
    try:
        return json.loads(response_text)
    except json.JSONDecodeError:
        json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response_text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group(1))
        start = response_text.find('{')
        end = response_text.rfind('}') + 1
        if start >= 0 and end > start:
            return json.loads(response_text[start:end])
        raise ValueError("Could not parse JSON from Claude response")


def _merge_extracted_chunks(partials: list[dict]) -> dict:
    """
    Merge extracted dicts from multiple transcript chunks.
    - Lists (decizii, ordine_de_zi, actiuni, participanti, voturi) are concatenated + de-duped.
    - Strings (concluzia, scurta_descriere, tematica) are joined with ' | '.
    - Scalars: the first non-empty wins.
    """
    if not partials:
        return {}
    if len(partials) == 1:
        return partials[0]

    merged: dict = {}
    LIST_KEYS = {
        "decizii", "decizii_luate", "hotarari", "ordine_de_zi", "subiecte_discutate",
        "agenda", "actiuni_de_urmat", "actions", "participanti", "participants",
        "voturi", "amenadamente", "intrebari",
    }

    all_keys = set()
    for p in partials:
        if isinstance(p, dict):
            all_keys.update(p.keys())

    for key in all_keys:
        values = [p.get(key) for p in partials if isinstance(p, dict) and p.get(key) is not None]
        if not values:
            continue

        if key in LIST_KEYS:
            combined: list = []
            seen: set = set()
            for v in values:
                if isinstance(v, list):
                    for item in v:
                        marker = json.dumps(item, ensure_ascii=False, sort_keys=True) if isinstance(item, (dict, list)) else str(item)
                        if marker not in seen:
                            seen.add(marker)
                            combined.append(item)
                elif v:
                    marker = str(v)
                    if marker not in seen:
                        seen.add(marker)
                        combined.append(v)
            merged[key] = combined
        elif all(isinstance(v, str) for v in values):
            non_empty = [v.strip() for v in values if v and v.strip()]
            if non_empty:
                merged[key] = non_empty[0] if len(non_empty) == 1 else " | ".join(dict.fromkeys(non_empty))
        else:
            merged[key] = values[0]

    return merged


async def extract_meeting_data(transcript: str, vertical_type: str = "GAL") -> dict:
    """
    Extract structured data from transcript using Anthropic Claude SDK.
    For long recordings (>12k chars) the transcript is chunked, each chunk is
    extracted individually, and the partial results are merged so NO content is lost
    (proces verbal must cover the COMPLETE recording — even 5h+ meetings).
    """
    from verticals import get_vertical_config

    vertical_config = get_vertical_config(vertical_type)
    CHUNK_SIZE = 12000

    # Short recording — single call (fast path)
    if len(transcript) <= CHUNK_SIZE:
        response = await anthropic_client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=4000,
            system=vertical_config.prompt_template,
            messages=[{
                "role": "user",
                "content": f"Extrage informațiile structurate din această transcriere:\n\n{transcript}"
            }]
        )
        return _parse_claude_json(response.content[0].text)

    # Long recording — chunk by sentence boundaries, extract partials, merge
    chunks: list[str] = []
    start = 0
    while start < len(transcript):
        end = start + CHUNK_SIZE
        if end < len(transcript):
            # Prefer to split on a paragraph/sentence boundary so context is preserved
            for sep in ("\n\n", "\n", ". ", " "):
                cut = transcript.rfind(sep, start, end)
                if cut > start + CHUNK_SIZE // 2:
                    end = cut + len(sep)
                    break
        chunks.append(transcript[start:end])
        start = end

    print(f"[Extract] Long transcript: {len(transcript)} chars → {len(chunks)} chunks (full coverage, no truncation)")

    chunk_system = vertical_config.prompt_template + (
        "\n\nIMPORTANT: Aceasta este O PARTE dintr-o transcriere mai lungă. "
        "Extrage TOATE informațiile prezente în această parte. NU omite nimic. "
        "Listele (decizii, ordine de zi, acțiuni, participanți) trebuie să conțină "
        "FIECARE element menționat în această parte. Returnează JSON valid."
    )

    partials: list[dict] = []
    # Limit concurrency to keep within Anthropic rate limits
    sem = asyncio.Semaphore(3)

    async def _extract_one(idx: int, chunk: str) -> dict:
        async with sem:
            try:
                resp = await anthropic_client.messages.create(
                    model="claude-sonnet-4-5",
                    max_tokens=4000,
                    system=chunk_system,
                    messages=[{
                        "role": "user",
                        "content": f"Partea {idx+1}/{len(chunks)} din transcriere:\n\n{chunk}"
                    }]
                )
                parsed = _parse_claude_json(resp.content[0].text)
                print(f"[Extract] Chunk {idx+1}/{len(chunks)} parsed ({len(parsed)} fields)")
                return parsed
            except Exception as exc:
                print(f"[Extract] Chunk {idx+1}/{len(chunks)} FAILED: {exc}")
                return {}

    partials = await asyncio.gather(*[_extract_one(i, c) for i, c in enumerate(chunks)])
    merged = _merge_extracted_chunks([p for p in partials if p])

    print(f"[Extract] Merged {len(partials)} chunks → {len(merged)} fields")
    return merged


async def process_meeting(meeting_id: str):
    """Background task: transcribe + extract data for a meeting."""
    try:
        meeting = await meetings_col.find_one({"_id": ObjectId(meeting_id)})
        if not meeting:
            return
        
        # Update status to processing
        await meetings_col.update_one(
            {"_id": ObjectId(meeting_id)},
            {"$set": {"status": "processing", "updated_at": datetime.now(timezone.utc)}}
        )
        
        audio_path = meeting.get("audio_path")
        if not audio_path or not os.path.exists(audio_path):
            await meetings_col.update_one(
                {"_id": ObjectId(meeting_id)},
                {"$set": {"status": "error", "error": "Fișier audio lipsă", "updated_at": datetime.now(timezone.utc)}}
            )
            return

        # Resolve vertical early — needed by transcription + diarization
        vertical_type = meeting.get("vertical_type") or "GENERAL"

        # Preprocess + always-split + parallel transcribe
        print(f"[Pipeline] Transcribing meeting {meeting_id} (vertical: {vertical_type})...")
        transcription = await transcribe_with_chunking(audio_path, vertical_type=vertical_type)

        # TASK 3 — Persist chunk processing metadata for production monitoring
        await meetings_col.update_one(
            {"_id": ObjectId(meeting_id)},
            {"$set": {
                "processing_chunks_total":     transcription.get("chunks_total", 1),
                "processing_chunks_succeeded": transcription.get("chunks_succeeded", 1),
            }}
        )

        # Auto-detect vertical when user left default "GENERAL"
        if vertical_type == "GENERAL" and transcription.get("text"):
            detected = await detect_vertical(transcription["text"][:1500])
            if detected != "GENERAL":
                print(f"[Pipeline] Auto-detected vertical: {detected}")
                vertical_type = detected
                await meetings_col.update_one(
                    {"_id": ObjectId(meeting_id)},
                    {"$set": {
                        "vertical_type":          vertical_type,
                        "vertical_auto_detected": True,
                    }}
                )

        # TASK 5 — solo_reading bypasses diarization entirely
        if vertical_type == "SOLO_READING":
            diarized = [{
                "speaker":   "Vorbitor",
                "role":      None,
                "timestamp": "00:00",
                "text":      transcription["text"],
            }]
            print("[Pipeline] Solo reading — skipping diarization")
        else:
            # Step 1c: Speaker diarization — vertical-aware two-pass NLP (TASK 4)
            print(f"[Pipeline] Diarizing speakers for meeting {meeting_id}...")
            diarized = await diarize_transcript(
                segments=transcription["segments"],
                words=transcription.get("words", []),
                audio_path=audio_path,
                vertical_type=vertical_type,
            )
            print(f"[Pipeline] Found {len(set(d.get('speaker','') for d in diarized))} speakers in {len(diarized)} segments")

        await meetings_col.update_one(
            {"_id": ObjectId(meeting_id)},
            {"$set": {
                "transcript":          transcription["text"],
                "segments":            transcription["segments"],
                "words":               transcription.get("words", []),
                "diarized_transcript": diarized,
                "updated_at":          datetime.now(timezone.utc),
            }}
        )

        # Step 2: Extract structured data via Claude
        print(f"[Meetings.ro] Extracting data for meeting {meeting_id} (vertical: {vertical_type})...")
        extracted = await extract_meeting_data(transcription["text"], vertical_type)
        print(f"[Meetings.ro] Extracted fields: {list(extracted.keys())}")

        # Determine locality from extracted fields — try multiple possible keys
        raw_locality = (
            extracted.get("localitate") or
            extracted.get("locality") or
            extracted.get("loc_desfasurare") or
            extracted.get("location") or
            ""
        )
        # Reject placeholder values
        INVALID_LOCALITY = {"necunoscut", "unknown", "n/a", "na", "null", "none", "-", ""}
        locality = raw_locality.strip() if raw_locality and raw_locality.strip().lower() not in INVALID_LOCALITY else None

        # Generate title: DD.MM.YYYY | Localitate  OR  DD.MM.YYYY | HH:MM
        created_at = meeting.get("created_at", datetime.now(timezone.utc))
        if isinstance(created_at, str):
            created_at = datetime.fromisoformat(created_at.replace("Z", "+00:00"))
        if locality:
            title = f"{created_at.strftime('%d.%m.%Y')} | {locality}"
        else:
            title = created_at.strftime("%d.%m.%Y | %H:%M")

        # Save locality as "Necunoscut" only for DB purposes (backward compat), not in title
        db_locality = locality or "Necunoscut"

        # Build update data — GAL fields at top level + vertical_config for all verticals
        update_data = {
            "title": title,
            "locality": db_locality,
            "vertical_config": extracted,
            "status": "done",
            "error": None,
            "updated_at": datetime.now(timezone.utc)
        }

        # Also save GAL-specific fields at top level for backward compatibility
        for field in ["data_desfasurare", "format_intalnire", "loc_desfasurare",
                       "mod_promovare", "obiectiv", "tematica", "scurta_descriere",
                       "numar_participanti", "concluzia"]:
            if field in extracted:
                update_data[field] = extracted[field]

        await meetings_col.update_one(
            {"_id": ObjectId(meeting_id)},
            {"$set": update_data}
        )

        # ---- Charge minutes against user's plan (status is now "done") ----
        try:
            duration_minutes = float(meeting.get("duration_seconds", 0.0) or 0.0) / 60.0
            owner_id = meeting.get("user_id")
            if duration_minutes > 0 and owner_id:
                await users_col.update_one(
                    {"_id": ObjectId(owner_id)},
                    {"$inc": {"minutes_used_this_month": duration_minutes}}
                )
                owner = await users_col.find_one({"_id": ObjectId(owner_id)})
                owner_plan = normalize_plan(owner.get("plan") if owner else None)
                print(f"[Usage] +{duration_minutes:.1f} min for user {owner_id} | plan: {owner_plan}")
        except Exception as usage_err:
            print(f"[Usage] Failed to increment minutes for meeting {meeting_id}: {usage_err}")

        # Ensure locality exists in localities collection (only real localities, not placeholders)
        if locality:
            await localities_col.update_one(
                {"name": locality},
                {"$setOnInsert": {"name": locality, "created_at": datetime.now(timezone.utc)}, "$inc": {"count": 1}},
                upsert=True
            )
        
        print(f"[GAL] Meeting {meeting_id} processed successfully!")

        # TASK 6 — Respect user's audio retention preference
        # Default (GDPR): delete after processing. Opt-in: keep on server.
        owner_user = None
        try:
            if meeting.get("user_id"):
                owner_user = await users_col.find_one({"_id": ObjectId(meeting["user_id"])})
        except Exception:
            owner_user = None
        keep_audio = bool(owner_user.get("keep_audio", False)) if owner_user else False

        if keep_audio:
            print(f"[Cleanup] Audio retained per user preference: {audio_path}")
            await meetings_col.update_one(
                {"_id": ObjectId(meeting_id)},
                {"$set": {"audio_retained": True}},
            )
        elif audio_path and os.path.exists(audio_path):
            try:
                os.remove(audio_path)
                await meetings_col.update_one(
                    {"_id": ObjectId(meeting_id)},
                    {"$set": {
                        "audio_path":       None,
                        "audio_url":        None,
                        "audio_deleted_at": datetime.now(timezone.utc),
                        "audio_retained":   False,
                    }}
                )
                print(f"[GDPR] Audio deleted: {audio_path}")
            except Exception as gdpr_err:
                print(f"[GDPR] Could not delete audio {audio_path}: {gdpr_err}")

        # Send push notification to user
        meeting_doc = await meetings_col.find_one({"_id": ObjectId(meeting_id)})
        if meeting_doc and meeting_doc.get("user_id"):
            await notify_user_meeting_done(meeting_doc["user_id"], title)

    except Exception as e:
        print(f"[GAL] Error processing meeting {meeting_id}: {e}")
        import traceback
        traceback.print_exc()
        await meetings_col.update_one(
            {"_id": ObjectId(meeting_id)},
            {"$set": {
                "status": "error",
                "error": str(e),
                "updated_at": datetime.now(timezone.utc)
            }}
        )


# ==================== API ENDPOINTS ====================

@app.get("/api/v1/verticals")
async def get_verticals():
    """Get all available vertical configurations with output fields."""
    from verticals import VERTICALS
    result = []
    for key, config in VERTICALS.items():
        result.append({
            "id": key,
            "name": config.name,
            "display_name_ro": config.display_name_ro,
            "icon": config.icon,
            "color_accent": config.color_accent,
            "description_ro": config.description_ro,
            "output_fields": [
                {"key": f.key, "label_ro": f.label_ro, "field_type": f.field_type}
                for f in config.output_fields
            ],
            "predefined_locations": config.predefined_locations or [],
        })
    return {"verticals": result}


@app.get("/download/meetings-ro.zip")
async def download_project():
    """Download Meetings.ro Expo project"""
    from fastapi.responses import FileResponse
    zip_path = "/tmp/meetings-ro.zip"
    if not os.path.exists(zip_path):
        raise HTTPException(status_code=404, detail="Project archive not found")
    return FileResponse(
        zip_path,
        media_type="application/zip",
        filename="meetings-ro.zip"
    )


@app.get("/api/health")
async def health():
    return {"status": "ok", "service": "Meetings.ro API"}


# ==================== TENANTS ====================

class TenantCreate(BaseModel):
    name: str
    type: str  # primarie | ong | firma
    vertical: str = "GENERAL"

class InviteRequest(BaseModel):
    email: str
    role: str = "member"

class RegisterWithInviteRequest(BaseModel):
    token: str
    name: str
    password: str


@app.post("/api/tenants")
async def create_tenant(body: TenantCreate, user: dict = Depends(get_current_user)):
    """Create a new tenant (organization). First user becomes admin."""
    existing = await tenants_col.find_one({"created_by": str(user["_id"])})
    if existing:
        raise HTTPException(status_code=400, detail="Ai deja o organizație creată")

    # Also block if already in a tenant (not created by him)
    if user.get("tenant_id"):
        raise HTTPException(status_code=400, detail="Ești deja membre al unei organizații")

    valid_types = {"primarie", "ong", "firma", "consiliu", "spital", "scoala", "other"}
    if body.type not in valid_types:
        raise HTTPException(status_code=400, detail=f"Tip invalid. Valori acceptate: {valid_types}")

    valid_verticals = {"GENERAL", "GAL", "BANKING", "LEGAL", "JOURNALISM", "HEALTHCARE", "STARTUPS"}
    vertical = body.vertical.upper() if body.vertical else "GENERAL"
    if vertical not in valid_verticals:
        vertical = "GENERAL"

    now = datetime.now(timezone.utc)
    tenant = {
        "name": body.name,
        "type": body.type,
        "vertical": vertical,
        "plan": "free",
        "billing_email": user["email"],
        "antet_text": body.name,
        "stema_url": None,
        "created_by": str(user["_id"]),
        "created_at": now,
    }
    result = await tenants_col.insert_one(tenant)
    tenant_id = str(result.inserted_id)

    # Promote creator to admin of the new tenant
    await users_col.update_one(
        {"_id": user["_id"]},
        {"$set": {"tenant_id": tenant_id, "role": "admin"}}
    )

    tenant["_id"] = tenant_id
    tenant["created_at"] = now.isoformat()
    return {"tenant_id": tenant_id, **{k: v for k, v in tenant.items() if k != "_id"}}


@app.get("/api/tenants/me")
async def get_my_tenant(user: dict = Depends(get_current_user)):
    """Get current user's tenant."""
    tenant_id = user.get("tenant_id")
    if not tenant_id:
        raise HTTPException(status_code=404, detail="Nu ești parte dintr-o organizație")
    tenant = await tenants_col.find_one({"_id": ObjectId(tenant_id)})
    if not tenant:
        raise HTTPException(status_code=404, detail="Organizația nu există")
    return serialize_doc(tenant)


@app.patch("/api/tenants/me")
async def update_my_tenant(
    name: Optional[str] = None,
    antet_text: Optional[str] = None,
    vertical: Optional[str] = None,
    user: dict = Depends(require_role("admin")),
):
    """Update tenant info (admin only)."""
    tenant_id = user.get("tenant_id")
    if not tenant_id:
        raise HTTPException(status_code=404, detail="Nu ești parte dintr-o organizație")
    updates: dict = {}
    if name:
        updates["name"] = name
    if antet_text:
        updates["antet_text"] = antet_text
    if vertical:
        updates["vertical"] = vertical.upper()
    if updates:
        await tenants_col.update_one({"_id": ObjectId(tenant_id)}, {"$set": updates})
    tenant = await tenants_col.find_one({"_id": ObjectId(tenant_id)})
    return serialize_doc(tenant)


@app.get("/api/tenants/me/members")
async def get_tenant_members(user: dict = Depends(get_current_user)):
    """List all members of the current user's tenant."""
    tenant_id = user.get("tenant_id")
    if not tenant_id:
        raise HTTPException(status_code=404, detail="Nu ești parte dintr-o organizație")
    cursor = users_col.find(
        {"tenant_id": tenant_id},
        {"password_hash": 0, "verify_token": 0, "reset_token": 0}
    )
    members = []
    async for m in cursor:
        members.append(serialize_doc(m))
    return {"members": members, "count": len(members)}


@app.delete("/api/tenants/me/members/{user_id}")
async def remove_tenant_member(user_id: str, user: dict = Depends(require_role("admin"))):
    """Remove a member from the tenant (admin only). Cannot remove yourself."""
    if user_id == str(user["_id"]):
        raise HTTPException(status_code=400, detail="Nu poți elimina propriul cont din organizație")
    target = await users_col.find_one({"_id": ObjectId(user_id)})
    if not target or target.get("tenant_id") != user.get("tenant_id"):
        raise HTTPException(status_code=404, detail="Membrul nu a fost găsit în organizație")
    await users_col.update_one(
        {"_id": ObjectId(user_id)},
        {"$set": {"tenant_id": None, "role": "member"}}
    )
    return {"message": "Membrul a fost eliminat din organizație"}


# ==================== INVITE FLOW ====================

INVITE_EMAIL_TEMPLATE = """
<div style="font-family:sans-serif;max-width:480px;margin:40px auto;padding:0 24px;background:#fff;border-radius:12px;border:1px solid #E5E7EB">
  <div style="padding:32px 0 24px">
    <h1 style="color:#1B2A4A;font-size:22px;margin:0 0 8px">Ai fost invitat în {tenant_name}</h1>
    <p style="color:#6B7280;font-size:15px;margin:0 0 24px">
      Rolul tău va fi: <strong style="color:#1B2A4A">{role_label}</strong>
    </p>
    <a href="{accept_url}"
       style="display:inline-block;background:#1B2A4A;color:#FAF8F3;padding:14px 28px;
              border-radius:8px;text-decoration:none;font-size:15px;font-weight:600">
        Acceptă invitația →
    </a>
    <p style="color:#9CA3AF;font-size:12px;margin-top:24px">Link valabil 7 zile. Dacă nu ești tu, ignoră acest email.</p>
  </div>
  <div style="border-top:1px solid #F3F4F6;padding:16px 0;text-align:center">
    <span style="color:#9CA3AF;font-size:12px">Meetings.ro — Transcriere AI pentru organizații</span>
  </div>
</div>
"""

ROLE_LABELS_RO = {
    "member": "Membru", "secretary": "Secretar", "mayor": "Primar",
    "councilor": "Consilier", "clerk": "Funcționar", "admin": "Administrator"
}
VALID_ROLES = {"member", "secretary", "mayor", "councilor", "clerk", "admin"}
API_BASE = os.environ.get("API_BASE_URL", "https://meetings-ro-api.onrender.com")


@app.post("/api/tenants/invite")
async def invite_member(body: InviteRequest, user: dict = Depends(require_role("admin"))):
    """Invite a user to the tenant by email (admin only)."""
    tenant_id = user.get("tenant_id")
    if not tenant_id:
        raise HTTPException(status_code=400, detail="Nu ești parte dintr-o organizație")

    role = body.role if body.role in VALID_ROLES else "member"
    email = body.email.lower().strip()

    # Check if already in this tenant
    existing = await users_col.find_one({"email": email, "tenant_id": tenant_id})
    if existing:
        raise HTTPException(status_code=400, detail="Utilizatorul este deja în organizație")

    # Revoke any old pending invite for same email+tenant
    await invitations_col.delete_many({"email": email, "tenant_id": tenant_id})

    invite_token = secrets.token_urlsafe(32)
    now = datetime.now(timezone.utc)
    invite = {
        "tenant_id": tenant_id,
        "email": email,
        "role": role,
        "token": invite_token,
        "invited_by": str(user["_id"]),
        "expires_at": now + timedelta(days=7),
        "created_at": now,
    }
    await invitations_col.insert_one(invite)

    tenant = await tenants_col.find_one({"_id": ObjectId(tenant_id)})
    tenant_name = tenant["name"] if tenant else "Meetings.ro"
    accept_url = f"{API_BASE}/api/auth/accept-invite?token={invite_token}"
    role_label = ROLE_LABELS_RO.get(role, role.capitalize())

    try:
        resend.Emails.send({
            "from": "Meetings.ro <noreply@resend.dev>",
            "to": [email],
            "subject": f"Invitație să te alături organizației {tenant_name}",
            "html": INVITE_EMAIL_TEMPLATE.format(
                tenant_name=tenant_name,
                role_label=role_label,
                accept_url=accept_url,
            ),
        })
    except Exception as e:
        print(f"[Invite] Email failed: {e}. Token: {invite_token}")
        # Don't fail — return token so admin can share manually if needed
        return {
            "message": f"Invitație creată (email eșuat). Token manual: {invite_token}",
            "token": invite_token,
            "accept_url": accept_url,
        }

    return {"message": f"Invitație trimisă la {email}", "accept_url": accept_url}


@app.get("/api/tenants/me/invitations")
async def list_pending_invitations(user: dict = Depends(require_role("admin"))):
    """List pending invitations for the current tenant (admin only)."""
    tenant_id = user.get("tenant_id")
    if not tenant_id:
        raise HTTPException(status_code=404, detail="Nu ești parte dintr-o organizație")
    now = datetime.now(timezone.utc)
    cursor = invitations_col.find({"tenant_id": tenant_id, "expires_at": {"$gt": now}})
    invites = []
    async for inv in cursor:
        inv["_id"] = str(inv["_id"])
        inv["expires_at"] = inv["expires_at"].isoformat()
        inv["created_at"] = inv["created_at"].isoformat()
        invites.append(inv)
    return {"invitations": invites}


@app.delete("/api/tenants/me/invitations/{token}")
async def revoke_invitation(token: str, user: dict = Depends(require_role("admin"))):
    """Revoke a pending invitation (admin only)."""
    tenant_id = user.get("tenant_id")
    result = await invitations_col.delete_one({"token": token, "tenant_id": tenant_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Invitație negăsită")
    return {"message": "Invitație revocată"}


@app.get("/api/auth/accept-invite")
async def accept_invite_page(token: str):
    """Landing page when user clicks the invite link."""
    invite = await invitations_col.find_one({
        "token": token,
        "expires_at": {"$gt": datetime.now(timezone.utc)}
    })
    if not invite:
        return HTMLResponse("""
            <html><body style="font-family:sans-serif;text-align:center;padding:60px;background:#FAF8F3">
            <h1 style="color:#EF4444">Link invalid sau expirat</h1>
            <p>Roagă administratorul să îți trimită o nouă invitație.</p>
            </body></html>
        """)

    existing_user = await users_col.find_one({"email": invite["email"]})
    if existing_user:
        # User already exists → just assign tenant and role
        await users_col.update_one(
            {"_id": existing_user["_id"]},
            {"$set": {"tenant_id": invite["tenant_id"], "role": invite["role"]}}
        )
        await invitations_col.delete_one({"token": token})
        return HTMLResponse("""
            <html><body style="font-family:sans-serif;text-align:center;padding:60px;background:#FAF8F3">
            <h1 style="color:#1B2A4A">✓ Ai fost adăugat în organizație!</h1>
            <p style="color:#6B7280">Deschide aplicația Meetings.ro și loghează-te din nou.</p>
            </body></html>
        """)

    # New user → show registration page
    tenant = await tenants_col.find_one({"_id": ObjectId(invite["tenant_id"])})
    tenant_name = tenant["name"] if tenant else "Meetings.ro"
    return HTMLResponse(f"""
        <html><head><meta name="viewport" content="width=device-width,initial-scale=1"></head>
        <body style="font-family:sans-serif;text-align:center;padding:60px;background:#FAF8F3;max-width:400px;margin:0 auto">
        <h1 style="color:#1B2A4A">Bun venit în {tenant_name}!</h1>
        <p style="color:#6B7280">Creează contul tău Meetings.ro folosind adresa<br>
           <strong style="color:#1B2A4A">{invite["email"]}</strong></p>
        <div style="background:#fff;border-radius:12px;padding:24px;border:1px solid #E5E7EB;margin-top:24px;text-align:left">
          <p style="color:#9CA3AF;font-size:12px;text-align:center">
            Folosește acest token în aplicație la înregistrare:<br>
            <code style="background:#F3F4F6;padding:4px 8px;border-radius:4px;font-size:11px">{token}</code>
          </p>
        </div>
        </body></html>
    """)


@app.post("/api/auth/register-with-invite")
async def register_with_invite(body: RegisterWithInviteRequest):
    """Register a new user via an invitation token."""
    invite = await invitations_col.find_one({
        "token": body.token,
        "expires_at": {"$gt": datetime.now(timezone.utc)}
    })
    if not invite:
        raise HTTPException(status_code=400, detail="Invitație invalidă sau expirată")

    # Check email not already taken
    existing = await users_col.find_one({"email": invite["email"]})
    if existing:
        # Already has account → just join tenant
        await users_col.update_one(
            {"_id": existing["_id"]},
            {"$set": {"tenant_id": invite["tenant_id"], "role": invite["role"]}}
        )
        await invitations_col.delete_one({"token": body.token})
        token_jwt = create_token(str(existing["_id"]), existing["email"])
        return {
            "token": token_jwt,
            "name": existing.get("name", ""),
            "email": existing["email"],
            "role": invite["role"],
            "tenant_id": invite["tenant_id"],
        }

    now = datetime.now(timezone.utc)
    new_user = {
        "email": invite["email"],
        "name": body.name,
        "password_hash": hash_password(body.password),
        "company": None,
        "plan": "FREE",
        "role": invite["role"],
        "tenant_id": invite["tenant_id"],
        # Invited users are pre-verified — they followed an org-issued link
        "email_verified": True,
        "minutes_used_this_month": 0.0,
        "last_monthly_reset": now,
        "created_at": now,
        "updated_at": now,
    }
    result = await users_col.insert_one(new_user)
    await invitations_col.delete_one({"token": body.token})

    token_jwt = create_token(str(result.inserted_id), invite["email"])
    return {
        "token": token_jwt,
        "name": body.name,
        "email": invite["email"],
        "role": invite["role"],
        "tenant_id": invite["tenant_id"],
    }


# ==================== ROLE MANAGEMENT ====================

class RoleUpdateBody(BaseModel):
    role: str


@app.patch("/api/users/{user_id}/role")
async def change_user_role(user_id: str, body: RoleUpdateBody, user: dict = Depends(require_role("admin"))):
    """Change the role of a tenant member (admin only)."""
    if body.role not in VALID_ROLES:
        raise HTTPException(status_code=400, detail=f"Rol invalid. Roluri valide: {list(VALID_ROLES)}")

    try:
        target = await users_col.find_one({"_id": ObjectId(user_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="ID user invalid")
    if not target:
        raise HTTPException(status_code=404, detail="Utilizator negăsit")
    if target.get("tenant_id") != user.get("tenant_id"):
        raise HTTPException(status_code=403, detail="Nu poți modifica utilizatori din altă organizație")
    if str(target["_id"]) == str(user["_id"]) and body.role != "admin":
        raise HTTPException(status_code=400, detail="Nu îți poți retrage propriul rol de admin")

    await users_col.update_one({"_id": ObjectId(user_id)}, {"$set": {"role": body.role}})
    return {"message": f"Rol actualizat la {ROLE_LABELS_RO.get(body.role, body.role)}", "role": body.role}


# ---- MEETINGS CRUD ----

@app.post("/api/meetings")
async def create_meeting(data: MeetingCreate = None, user: dict = Depends(get_current_user)):
    """Create a new meeting placeholder. Plan-quota enforcement happens at
    upload time (minute-based) rather than per-meeting."""
    # Email verification gate
    if not is_email_verified(user):
        raise HTTPException(
            status_code=403,
            detail={
                "error": "email_not_verified",
                "message": "Te rugăm să îți verifici adresa de email înainte de a folosi aplicația.",
            }
        )

    now = datetime.now(timezone.utc)

    # Determine vertical: explicit request > tenant config > GENERAL default
    vertical_type = data.vertical_type if data and hasattr(data, 'vertical_type') and data.vertical_type else None
    if not vertical_type:
        if user.get("tenant_id"):
            tenant_doc = await tenants_col.find_one({"_id": ObjectId(user["tenant_id"])})
            vertical_type = tenant_doc.get("vertical") if tenant_doc else "GENERAL"
        vertical_type = vertical_type or "GENERAL"

    meeting = {
        "title": (data.title if data and data.title else None),
        "locality": (data.locality if data and data.locality else None),
        "date": now.isoformat(),
        "audio_path": None,
        "audio_url": None,
        "transcript": None,
        "segments": [],
        # GAL report fields (backward compat)
        "data_desfasurare": None,
        "format_intalnire": None,
        "loc_desfasurare": None,
        "mod_promovare": None,
        "obiectiv": None,
        "tematica": None,
        "scurta_descriere": None,
        "numar_participanti": None,
        "concluzia": None,
        # Vertical system
        "vertical_type": vertical_type,
        "vertical_config": {},
        # Ownership + multi-tenant
        "user_id": str(user["_id"]),
        "tenant_id": str(user["tenant_id"]) if user.get("tenant_id") else None,
        "status": "pending",
        "error": None,
        "duration": 0,
        "created_at": now,
        "updated_at": now
    }
    result = await meetings_col.insert_one(meeting)
    meeting["_id"] = result.inserted_id

    return serialize_doc(meeting)


ALLOWED_AUDIO_TYPES = {
    "audio/mpeg", "audio/mp3", "audio/wav", "audio/x-wav",
    "audio/m4a", "audio/x-m4a", "audio/mp4", "audio/ogg",
    "audio/webm", "audio/aac", "audio/x-m4a", "audio/3gpp",
    "audio/x-aac", "audio/flac", "audio/x-flac",
    "application/octet-stream",  # iOS often sends this for m4a
    "video/mp4",  # some devices send m4a as video/mp4
}
ALLOWED_AUDIO_EXTENSIONS = {".m4a", ".mp3", ".wav", ".ogg", ".webm", ".aac", ".flac", ".3gp", ".mp4"}
MAX_FILE_SIZE = 200 * 1024 * 1024  # 200MB hard limit — backend will chunk anything >20MB for Groq

@app.post("/api/meetings/{meeting_id}/upload")
async def upload_audio(meeting_id: str, background_tasks: BackgroundTasks, file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    """Upload audio file for a meeting and start processing."""
    # ---- Email verification gate ----
    if not is_email_verified(user):
        raise HTTPException(
            status_code=403,
            detail={
                "error": "email_not_verified",
                "message": "Te rugăm să îți verifici adresa de email înainte de a folosi aplicația.",
            }
        )

    # ---- Plan quota gate (minute-based) ----
    user = await reset_monthly_if_needed(user)
    plan_key, used_minutes, limit_minutes = get_plan_quota(user)
    if used_minutes >= limit_minutes:
        raise HTTPException(
            status_code=403,
            detail={
                "error": "plan_limit_reached",
                "message": (
                    f"Ai folosit {used_minutes:.0f} din {int(limit_minutes)} minute "
                    f"disponibile în planul tău {plan_key.capitalize()}. "
                    f"Fă upgrade pentru a continua."
                ),
                "minutes_used": round(used_minutes, 1),
                "minutes_limit": int(limit_minutes),
                "remaining_minutes": 0,
                "upgrade_required": True,
            }
        )

    # Validate MIME type — allow unknown types if extension is valid
    file_ext = Path(file.filename or "").suffix.lower() if file.filename else ""
    content_type = (file.content_type or "").lower().strip()

    if content_type and content_type not in ALLOWED_AUDIO_TYPES:
        # MIME unknown, check extension as fallback
        if file_ext not in ALLOWED_AUDIO_EXTENSIONS:
            print(f"[Upload] Rejected: content_type={content_type}, ext={file_ext}, filename={file.filename}")
            raise HTTPException(status_code=415, detail=f"Format audio neacceptat: {content_type}")

    # Validate meeting exists + ownership
    meeting = await verify_meeting_ownership(meeting_id, user)

    # Save audio file with size check
    ext = file.filename.split(".")[-1] if "." in file.filename else "webm"
    allowed_exts = {"mp3", "wav", "m4a", "mp4", "ogg", "webm", "aac", "3gpp"}
    if ext.lower() not in allowed_exts:
        ext = "webm"
    audio_filename = f"{meeting_id}.{ext}"
    audio_path = str(UPLOAD_DIR / audio_filename)

    content = await file.read()
    file_size_bytes = len(content)
    file_size_mb = file_size_bytes / (1024 * 1024)
    print(f"[Upload] Received file: {file.filename}, size={file_size_mb:.1f}MB, type={file.content_type}")

    if file_size_bytes > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"Fișierul depășește limita maximă de 200MB ({file_size_mb:.1f}MB primit). "
                   f"Comprimați înregistrarea sau împărțiți-o în segmente."
        )

    async with aiofiles.open(audio_path, "wb") as f:
        await f.write(content)

    file_size = os.path.getsize(audio_path)

    # Detect audio duration so we can charge minutes against the user's plan
    duration_seconds = await get_audio_duration_seconds(audio_path)
    print(f"[Upload] Duration detected: {duration_seconds:.1f}s ({duration_seconds/60:.1f} min)")

    # Update meeting with audio info
    await meetings_col.update_one(
        {"_id": ObjectId(meeting_id)},
        {"$set": {
            "audio_path": audio_path,
            "audio_url": f"/api/meetings/{meeting_id}/audio",
            "status": "uploading",
            "file_size": file_size,
            "duration_seconds": duration_seconds,
            "updated_at": datetime.now(timezone.utc)
        }}
    )
    
    # Start processing in background
    background_tasks.add_task(process_meeting, meeting_id)
    
    updated = await meetings_col.find_one({"_id": ObjectId(meeting_id)})
    return serialize_doc(updated)


@app.get("/api/meetings/{meeting_id}/audio")
async def get_audio(meeting_id: str, request: Request, user: dict = Depends(get_current_user)):
    """Stream audio file with HTTP Range request support for mobile playback."""
    meeting = await verify_meeting_ownership(meeting_id, user)
    if not meeting.get("audio_path"):
        raise HTTPException(status_code=404, detail="Audio nu a fost găsit")
    
    audio_path = meeting["audio_path"]
    if not os.path.exists(audio_path):
        raise HTTPException(status_code=404, detail="Fișier audio lipsă")
    
    ext = audio_path.split(".")[-1]
    media_types = {
        "webm": "audio/webm",
        "mp3": "audio/mpeg",
        "wav": "audio/wav",
        "m4a": "audio/mp4",
        "ogg": "audio/ogg"
    }
    media_type = media_types.get(ext, "audio/webm")
    file_size = os.path.getsize(audio_path)
    
    # Check for Range header
    range_header = None
    if request and request.headers.get("range"):
        range_header = request.headers.get("range")
    
    if range_header:
        # Parse range: "bytes=start-end"
        try:
            range_spec = range_header.replace("bytes=", "")
            parts = range_spec.split("-")
            start = int(parts[0]) if parts[0] else 0
            end = int(parts[1]) if parts[1] else file_size - 1
        except (ValueError, IndexError):
            start = 0
            end = file_size - 1
        
        # Clamp values
        start = max(0, min(start, file_size - 1))
        end = max(start, min(end, file_size - 1))
        content_length = end - start + 1
        
        async def range_file_generator():
            async with aiofiles.open(audio_path, "rb") as f:
                await f.seek(start)
                remaining = content_length
                while remaining > 0:
                    chunk_size = min(65536, remaining)
                    chunk = await f.read(chunk_size)
                    if not chunk:
                        break
                    remaining -= len(chunk)
                    yield chunk
        
        return StreamingResponse(
            range_file_generator(),
            status_code=206,
            media_type=media_type,
            headers={
                "Content-Range": f"bytes {start}-{end}/{file_size}",
                "Accept-Ranges": "bytes",
                "Content-Length": str(content_length),
                "Cache-Control": "public, max-age=3600",
            }
        )
    else:
        # Full file response
        async def full_file_generator():
            async with aiofiles.open(audio_path, "rb") as f:
                while chunk := await f.read(65536):
                    yield chunk
        
        return StreamingResponse(
            full_file_generator(),
            media_type=media_type,
            headers={
                "Accept-Ranges": "bytes",
                "Content-Length": str(file_size),
                "Cache-Control": "public, max-age=3600",
            }
        )


@app.get("/api/meetings")
async def list_meetings(
    locality: Optional[str] = None,
    status: Optional[str] = None,
    q: Optional[str] = None,
    page: int = Query(1, ge=1),
    limit: int = Query(50, ge=1, le=100),
    user: dict = Depends(get_current_user),
):
    """List meetings with optional filters (scoped to current user / tenant)."""
    query = build_meetings_scope_query(user)

    if locality:
        query["locality"] = locality
    if status:
        query["status"] = status
    if q:
        query["$text"] = {"$search": q}
    
    skip = (page - 1) * limit
    
    total = await meetings_col.count_documents(query)
    cursor = meetings_col.find(query).sort("created_at", -1).skip(skip).limit(limit)
    meetings = []
    async for doc in cursor:
        meetings.append(serialize_doc(doc))
    
    return {
        "meetings": meetings,
        "total": total,
        "page": page,
        "limit": limit,
        "pages": (total + limit - 1) // limit
    }


# ---- CALENDAR: MEETING DATES ----

@app.get("/api/meetings/calendar-dates")
async def get_meeting_dates(
    year: int = Query(...),
    month: int = Query(...),
    user: dict = Depends(get_current_user),
):
    """Get dates that have meetings for a given month (for calendar highlighting)."""
    from calendar import monthrange

    start_date = datetime(year, month, 1, tzinfo=timezone.utc)
    _, last_day = monthrange(year, month)
    end_date = datetime(year, month, last_day, 23, 59, 59, tzinfo=timezone.utc)

    scope = build_meetings_scope_query(user)
    match_stage = {**scope, "created_at": {"$gte": start_date, "$lte": end_date}}
    pipeline = [
        {"$match": match_stage},
        {"$group": {
            "_id": {"$dateToString": {"format": "%Y-%m-%d", "date": "$created_at"}},
            "count": {"$sum": 1}
        }},
        {"$sort": {"_id": 1}}
    ]
    
    cursor = meetings_col.aggregate(pipeline)
    dates = {}
    async for doc in cursor:
        dates[doc["_id"]] = doc["count"]
    
    return {"dates": dates}


@app.get("/api/meetings/calendar-by-date")
async def get_meetings_by_date(
    date: str = Query(...),
    user: dict = Depends(get_current_user),
):
    """Get meetings for a specific date (YYYY-MM-DD)."""
    try:
        target_date = datetime.strptime(date, "%Y-%m-%d").replace(tzinfo=timezone.utc)
    except ValueError:
        raise HTTPException(status_code=400, detail="Format dată invalid. Folosiți YYYY-MM-DD")

    next_date = target_date.replace(hour=23, minute=59, second=59)

    query = {
        **build_meetings_scope_query(user),
        "created_at": {"$gte": target_date, "$lte": next_date}
    }
    
    cursor = meetings_col.find(query).sort("created_at", -1)
    meetings = []
    async for doc in cursor:
        meetings.append(serialize_doc(doc))
    
    return {"meetings": meetings, "date": date, "count": len(meetings)}


@app.get("/api/meetings/{meeting_id}")
async def get_meeting(meeting_id: str, user: dict = Depends(get_current_user)):
    """Get single meeting detail."""
    meeting = await verify_meeting_ownership(meeting_id, user)
    return serialize_doc(meeting)


@app.patch("/api/meetings/{meeting_id}")
async def update_meeting(meeting_id: str, data: MeetingUpdate, user: dict = Depends(get_current_user)):
    """Update meeting title or locality."""
    meeting = await verify_meeting_ownership(meeting_id, user)
    
    update_fields = {"updated_at": datetime.now(timezone.utc)}
    if data.title is not None:
        update_fields["title"] = data.title
    if data.locality is not None:
        update_fields["locality"] = data.locality
        # Ensure new locality exists
        if data.locality and data.locality != "Necunoscut":
            await localities_col.update_one(
                {"name": data.locality},
                {"$setOnInsert": {"name": data.locality, "created_at": datetime.now(timezone.utc)}, "$inc": {"count": 1}},
                upsert=True
            )
    
    await meetings_col.update_one(
        {"_id": ObjectId(meeting_id)},
        {"$set": update_fields}
    )
    
    updated = await meetings_col.find_one({"_id": ObjectId(meeting_id)})
    return serialize_doc(updated)


@app.delete("/api/meetings/{meeting_id}")
async def delete_meeting(meeting_id: str, user: dict = Depends(get_current_user)):
    """Delete a meeting."""
    meeting = await verify_meeting_ownership(meeting_id, user)
    
    # Delete audio file if exists
    if meeting.get("audio_path") and os.path.exists(meeting["audio_path"]):
        os.remove(meeting["audio_path"])
    
    await meetings_col.delete_one({"_id": ObjectId(meeting_id)})
    return {"status": "deleted"}


# ---- FEEDBACK ----

class FeedbackPayload(BaseModel):
    rating: int                          # 1 = positive, -1 = negative
    issues: Optional[List[str]] = []     # only on negative: selected issue chips
    comment: Optional[str] = None        # optional free text (future use)


@app.post("/api/meetings/{meeting_id}/feedback")
async def submit_feedback(
    meeting_id: str,
    payload: FeedbackPayload,
    user: dict = Depends(get_current_user),
):
    """
    Saves user quality feedback on a transcription.
    Used for production quality monitoring — not shown to other users.
    """
    if not ObjectId.is_valid(meeting_id):
        raise HTTPException(status_code=400, detail="Invalid meeting ID")

    meeting = await meetings_col.find_one({
        "_id": ObjectId(meeting_id),
        "user_id": user["_id"],
    })
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")

    feedback_doc = {
        "rating":           payload.rating,
        "issues":           payload.issues or [],
        "comment":          payload.comment,
        "submitted_at":     datetime.now(timezone.utc),
        "vertical_type":    meeting.get("vertical_type", "GENERAL"),
        "chunks_total":     meeting.get("processing_chunks_total", 1),
        "duration_seconds": meeting.get("duration"),
    }

    await meetings_col.update_one(
        {"_id": ObjectId(meeting_id)},
        {"$set": {"feedback": feedback_doc}},
    )

    rating_label = "👍 POSITIVE" if payload.rating == 1 else "👎 NEGATIVE"
    print(
        f"[Feedback] {rating_label} | meeting={meeting_id} | "
        f"issues={payload.issues} | vertical={meeting.get('vertical_type')}"
    )

    return {"ok": True}


@app.get("/api/meetings/{meeting_id}/feedback")
async def get_feedback(
    meeting_id: str,
    user: dict = Depends(get_current_user),
):
    """Returns existing feedback for a meeting, if any."""
    if not ObjectId.is_valid(meeting_id):
        raise HTTPException(status_code=400, detail="Invalid meeting ID")

    meeting = await meetings_col.find_one(
        {"_id": ObjectId(meeting_id), "user_id": user["_id"]},
        {"feedback": 1},
    )
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")

    return meeting.get("feedback", None)


# ---- ACTION ITEMS ----

@app.patch("/api/meetings/{meeting_id}/actions/{action_id}")
async def toggle_action(meeting_id: str, action_id: str, user: dict = Depends(get_current_user)):
    """Toggle action item completion."""
    meeting = await verify_meeting_ownership(meeting_id, user)
    
    actions = meeting.get("actions", [])
    found = False
    for action in actions:
        if action["id"] == action_id:
            action["completed"] = not action["completed"]
            found = True
            break
    
    if not found:
        raise HTTPException(status_code=404, detail="Acțiunea nu a fost găsită")
    
    await meetings_col.update_one(
        {"_id": ObjectId(meeting_id)},
        {"$set": {"actions": actions, "updated_at": datetime.now(timezone.utc)}}
    )
    
    updated = await meetings_col.find_one({"_id": ObjectId(meeting_id)})
    return serialize_doc(updated)


# ---- REGENERATE AI ----

@app.post("/api/meetings/{meeting_id}/regenerate")
async def regenerate_meeting(meeting_id: str, background_tasks: BackgroundTasks, user: dict = Depends(get_current_user)):
    """Regenerate AI processing for a meeting."""
    meeting = await verify_meeting_ownership(meeting_id, user)
    
    if not meeting.get("transcript") and not meeting.get("audio_path"):
        raise HTTPException(status_code=400, detail="Nu există transcriere sau audio pentru procesare")
    
    # Reset status
    await meetings_col.update_one(
        {"_id": ObjectId(meeting_id)},
        {"$set": {"status": "processing", "error": None, "updated_at": datetime.now(timezone.utc)}}
    )
    
    if meeting.get("transcript"):
        # Just re-process the transcript
        background_tasks.add_task(reprocess_transcript, meeting_id)
    else:
        # Full reprocess from audio
        background_tasks.add_task(process_meeting, meeting_id)
    
    return {"status": "regenerating"}


# ---- CORRECT TRANSCRIPT (Claude) ----

@app.post("/api/meetings/{meeting_id}/correct-transcript")
async def correct_transcript(meeting_id: str, background_tasks: BackgroundTasks, user: dict = Depends(get_current_user)):
    """Use Claude to correct grammar/transcription errors, then re-extract report."""
    meeting = await verify_meeting_ownership(meeting_id, user)

    if not meeting.get("transcript"):
        raise HTTPException(status_code=400, detail="Nu există transcriere de corectat")

    # Save original transcript before correction
    original = meeting.get("transcript_original") or meeting.get("transcript")
    await meetings_col.update_one(
        {"_id": ObjectId(meeting_id)},
        {"$set": {
            "transcript_original": original,
            "status": "processing",
            "error": None,
            "updated_at": datetime.now(timezone.utc),
        }}
    )

    background_tasks.add_task(correct_and_reprocess, meeting_id)
    return {"status": "correcting"}


async def correct_and_reprocess(meeting_id: str):
    """Background: correct transcript with Claude, then re-extract report."""
    try:
        meeting = await meetings_col.find_one({"_id": ObjectId(meeting_id)})
        if not meeting or not meeting.get("transcript"):
            return

        # Step 1: Correct transcript with Claude
        response = await anthropic_client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=8000,
            system=(
                "Ești un editor profesionist de transcrieri în limba română. "
                "Corectează erorile gramaticale, de ortografie și de transcriere. "
                "Păstrează sensul original și structura textului. "
                "NU adăuga, nu șterge și nu reformula conținutul — doar corectează greșelile. "
                "Returnează DOAR textul corectat, fără explicații."
            ),
            messages=[{
                "role": "user",
                "content": f"Corectează următoarea transcriere:\n\n{meeting['transcript']}"
            }]
        )

        corrected_transcript = response.content[0].text.strip()

        # Step 2: Save corrected transcript
        await meetings_col.update_one(
            {"_id": ObjectId(meeting_id)},
            {"$set": {
                "transcript": corrected_transcript,
                "transcript_corrected_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc),
            }}
        )

        # Step 3: Re-extract report from corrected transcript
        extracted = await extract_meeting_data(corrected_transcript)

        locality = extracted.get("locality") or meeting.get("locality") or "Necunoscut"
        created_at = meeting.get("created_at", datetime.now(timezone.utc))
        if isinstance(created_at, str):
            created_at = datetime.fromisoformat(created_at.replace("Z", "+00:00"))
        title = f"{created_at.strftime('%d.%m.%Y')} | {locality}"

        await meetings_col.update_one(
            {"_id": ObjectId(meeting_id)},
            {"$set": {
                "title": title,
                "locality": locality,
                "data_desfasurare": extracted.get("data_desfasurare"),
                "format_intalnire": extracted.get("format_intalnire"),
                "loc_desfasurare": extracted.get("loc_desfasurare"),
                "mod_promovare": extracted.get("mod_promovare"),
                "obiectiv": extracted.get("obiectiv"),
                "tematica": extracted.get("tematica"),
                "scurta_descriere": extracted.get("scurta_descriere"),
                "numar_participanti": extracted.get("numar_participanti"),
                "concluzia": extracted.get("concluzia"),
                "status": "done",
                "error": None,
                "updated_at": datetime.now(timezone.utc),
            }}
        )

        if locality and locality != "Necunoscut":
            await localities_col.update_one(
                {"name": locality},
                {"$setOnInsert": {"name": locality, "created_at": datetime.now(timezone.utc)}, "$inc": {"count": 1}},
                upsert=True
            )
    except Exception as e:
        print(f"[CORRECT] Error correcting transcript {meeting_id}: {e}")
        await meetings_col.update_one(
            {"_id": ObjectId(meeting_id)},
            {"$set": {"status": "error", "error": str(e), "updated_at": datetime.now(timezone.utc)}}
        )


async def reprocess_transcript(meeting_id: str):
    """Re-extract data from existing transcript."""
    try:
        meeting = await meetings_col.find_one({"_id": ObjectId(meeting_id)})
        if not meeting or not meeting.get("transcript"):
            return
        
        extracted = await extract_meeting_data(meeting["transcript"])
        
        # Determine locality
        locality = extracted.get("locality") or meeting.get("locality") or "Necunoscut"
        
        # Generate title: DD.MM.YYYY | Localitatea
        created_at = meeting.get("created_at", datetime.now(timezone.utc))
        if isinstance(created_at, str):
            created_at = datetime.fromisoformat(created_at.replace("Z", "+00:00"))
        title = f"{created_at.strftime('%d.%m.%Y')} | {locality}"
        
        await meetings_col.update_one(
            {"_id": ObjectId(meeting_id)},
            {"$set": {
                "title": title,
                "locality": locality,
                "data_desfasurare": extracted.get("data_desfasurare"),
                "format_intalnire": extracted.get("format_intalnire"),
                "loc_desfasurare": extracted.get("loc_desfasurare"),
                "mod_promovare": extracted.get("mod_promovare"),
                "obiectiv": extracted.get("obiectiv"),
                "tematica": extracted.get("tematica"),
                "scurta_descriere": extracted.get("scurta_descriere"),
                "numar_participanti": extracted.get("numar_participanti"),
                "concluzia": extracted.get("concluzia"),
                "status": "done",
                "error": None,
                "updated_at": datetime.now(timezone.utc)
            }}
        )
        
        if locality and locality != "Necunoscut":
            await localities_col.update_one(
                {"name": locality},
                {"$setOnInsert": {"name": locality, "created_at": datetime.now(timezone.utc)}, "$inc": {"count": 1}},
                upsert=True
            )
    except Exception as e:
        print(f"[GAL] Error reprocessing {meeting_id}: {e}")
        await meetings_col.update_one(
            {"_id": ObjectId(meeting_id)},
            {"$set": {"status": "error", "error": str(e), "updated_at": datetime.now(timezone.utc)}}
        )


# ---- LOCALITIES ----

class LocalityCreate(BaseModel):
    name: str

class LocalityRename(BaseModel):
    new_name: str

@app.get("/api/localities")
async def list_localities(user: dict = Depends(get_current_user)):
    """List all localities (folders) with meeting counts."""
    # Get all localities from the localities collection
    all_localities = {}
    cursor = localities_col.find({}).sort("name", 1)
    async for doc in cursor:
        all_localities[doc["name"]] = {"name": doc["name"], "count": 0, "is_default": doc.get("is_default", False)}
    
    # Count meetings per locality
    pipeline = [
        {"$match": {"locality": {"$ne": None, "$ne": "", "$exists": True}}},
        {"$group": {"_id": "$locality", "count": {"$sum": 1}}},
        {"$sort": {"_id": 1}}
    ]
    agg_cursor = meetings_col.aggregate(pipeline)
    async for doc in agg_cursor:
        name = doc["_id"]
        if not name:
            continue
        if name in all_localities:
            all_localities[name]["count"] = doc["count"]
        else:
            all_localities[name] = {"name": name, "count": doc["count"], "is_default": False}
    
    localities = sorted(all_localities.values(), key=lambda x: x["name"] or "")
    return {"localities": localities}


@app.post("/api/localities")
async def create_locality(data: LocalityCreate, user: dict = Depends(get_current_user)):
    """Create a new locality folder."""
    name = data.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="Numele localității nu poate fi gol")
    
    # Check if exists
    existing = await localities_col.find_one({"name": name})
    if existing:
        raise HTTPException(status_code=409, detail="Localitate deja existentă")
    
    await localities_col.insert_one({
        "name": name,
        "created_at": datetime.now(timezone.utc),
        "is_default": False
    })
    
    return {"name": name, "count": 0, "is_default": False}


@app.patch("/api/localities/{locality_name}")
async def rename_locality(locality_name: str, data: LocalityRename, user: dict = Depends(get_current_user)):
    """Rename a locality folder and update all meetings."""
    new_name = data.new_name.strip()
    if not new_name:
        raise HTTPException(status_code=400, detail="Numele nou nu poate fi gol")
    
    # Check if source exists
    existing = await localities_col.find_one({"name": locality_name})
    if not existing:
        raise HTTPException(status_code=404, detail="Localitatea nu a fost găsită")
    
    # Check if target name already exists
    if new_name != locality_name:
        target = await localities_col.find_one({"name": new_name})
        if target:
            raise HTTPException(status_code=409, detail="O localitate cu acest nume există deja")
    
    # Rename in localities collection
    await localities_col.update_one(
        {"name": locality_name},
        {"$set": {"name": new_name, "updated_at": datetime.now(timezone.utc)}}
    )
    
    # Update all meetings with this locality
    result = await meetings_col.update_many(
        {"locality": locality_name},
        {"$set": {"locality": new_name}}
    )
    
    return {"old_name": locality_name, "new_name": new_name, "meetings_updated": result.modified_count}


@app.delete("/api/localities/{locality_name}")
async def delete_locality(locality_name: str, user: dict = Depends(get_current_user)):
    """Delete a locality folder (meetings stay but lose locality)."""
    # Delete from localities collection if exists
    await localities_col.delete_one({"name": locality_name})
    
    # Set meetings locality to null
    await meetings_col.update_many(
        {"locality": locality_name},
        {"$set": {"locality": None}}
    )
    
    return {"status": "deleted", "name": locality_name}




# ---- EXPORT ----

@app.get("/api/meetings/{meeting_id}/export/pdf")
async def export_pdf(meeting_id: str, user: dict = Depends(get_current_user)):
    """Export meeting as PDF."""
    meeting = await verify_meeting_ownership(meeting_id, user)
    
    from fpdf import FPDF
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)
    
    # Use built-in font that supports basic characters
    pdf.set_font("Arial", "B", 16)
    title = meeting.get("title", "Sedinta") or "Sedinta"
    # Transliterate Romanian chars for PDF compatibility
    title_safe = transliterate_ro(title)
    pdf.cell(0, 10, title_safe, new_x="LMARGIN", new_y="NEXT")
    
    pdf.set_font("Arial", "", 10)
    locality = meeting.get("locality", "Necunoscut") or "Necunoscut"
    date_str = meeting.get("date", "")
    if date_str:
        try:
            dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
            date_str = dt.strftime("%d.%m.%Y %H:%M")
        except (ValueError, AttributeError):
            pass
    
    pdf.cell(0, 6, f"Localitate: {transliterate_ro(locality)}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Data: {date_str}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    
    # Summary
    summary = meeting.get("summary", [])
    if summary:
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Rezumat", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Arial", "", 9)
        for i, item in enumerate(summary[:5]):  # Limit to 5 items
            text = transliterate_ro(item)[:100]  # Limit length
            pdf.cell(0, 5, f"{i+1}. {text}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)
    
    # Actions
    actions = meeting.get("actions", [])
    if actions:
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Actiuni", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Arial", "", 9)
        for i, action in enumerate(actions[:5]):  # Limit to 5 actions
            status = "[X]" if action.get("completed") else "[ ]"
            text = transliterate_ro(action.get("text", ""))[:80]
            pdf.cell(0, 5, f"{status} {text}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)
    
    # Output
    pdf_bytes = pdf.output()
    buffer = BytesIO(pdf_bytes)
    
    safe_title = re.sub(r'[^\w\s-]', '', title_safe).strip().replace(' ', '_')[:30]
    safe_date = date_str.replace('.', '-').replace(' ', '_').replace(':', '-') if date_str else 'no-date'
    filename = f"GAL_{safe_title}_{safe_date}.pdf"
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@app.get("/api/meetings/{meeting_id}/export/docx")
async def export_docx(meeting_id: str, user: dict = Depends(get_current_user)):
    """Export meeting as DOCX."""
    meeting = await verify_meeting_ownership(meeting_id, user)
    
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    title = meeting.get("title", "Ședință") or "Ședință"
    doc.add_heading(title, level=1)
    
    locality = meeting.get("locality", "Necunoscut") or "Necunoscut"
    date_str = meeting.get("date", "")
    if date_str:
        try:
            dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
            date_str = dt.strftime("%d.%m.%Y %H:%M")
        except (ValueError, AttributeError):
            pass
    
    p = doc.add_paragraph()
    p.add_run(f"Localitate: ").bold = True
    p.add_run(locality)
    p = doc.add_paragraph()
    p.add_run(f"Data: ").bold = True
    p.add_run(date_str)
    
    # Summary
    summary = meeting.get("summary", [])
    if summary:
        doc.add_heading("Rezumat", level=2)
        for item in summary:
            doc.add_paragraph(item, style='List Bullet')
    
    # Key Points
    key_points = meeting.get("key_points", [])
    if key_points:
        doc.add_heading("Puncte cheie", level=2)
        for item in key_points:
            doc.add_paragraph(item, style='List Bullet')
    
    # Actions
    actions = meeting.get("actions", [])
    if actions:
        doc.add_heading("Acțiuni", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Status'
        hdr_cells[1].text = 'Acțiune'
        hdr_cells[2].text = 'Responsabil'
        hdr_cells[3].text = 'Termen'
        
        for action in actions:
            row_cells = table.add_row().cells
            row_cells[0].text = '✓' if action.get('completed') else '○'
            row_cells[1].text = action.get('text', '')
            row_cells[2].text = action.get('owner', '-') or '-'
            row_cells[3].text = action.get('deadline', '-') or '-'
    
    # Transcript
    transcript = meeting.get("transcript", "")
    if transcript:
        doc.add_heading("Transcriere", level=2)
        doc.add_paragraph(transcript)
    
    # Output
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    safe_title = re.sub(r'[^\w\s-]', '', transliterate_ro(title)).strip().replace(' ', '_')[:50]
    filename = f"Sedinta_{safe_title}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ==================== PROCES VERBAL (Enterprise / Public Institutions) ====================

@app.get("/api/meetings/{meeting_id}/export/proces-verbal")
async def export_proces_verbal(meeting_id: str, user: dict = Depends(get_current_user)):
    """Export meeting as formal Proces Verbal (minutes) DOCX — institutional format.

    Used by mayors, secretaries, councilors and clerks in public institutions.
    Produces a formally structured document with:
      - Header (institution, date, participants, quorum)
      - Agenda (ordinea de zi)
      - Dezbateri (per speaker, from diarized transcript)
      - Hotărâri / Decizii adoptate
      - Semnături
    """
    meeting = await verify_meeting_ownership(meeting_id, user)

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Base style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ---- ANTET (Header) ----
    institution = meeting.get("locality") or user.get("company") or "Instituție"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(institution.upper())
    run.bold = True
    run.font.size = Pt(14)

    # Titlu
    title_text = meeting.get("title") or "Ședință"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nPROCES-VERBAL")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"al ședinței „{title_text}”").italic = True

    doc.add_paragraph()  # spacer

    # ---- DATA / LOC ----
    date_str = meeting.get("date") or meeting.get("data_desfasurare") or ""
    if date_str:
        try:
            dt = datetime.fromisoformat(str(date_str).replace("Z", "+00:00"))
            date_fmt = dt.strftime("%d.%m.%Y, ora %H:%M")
        except (ValueError, AttributeError):
            date_fmt = str(date_str)
    else:
        date_fmt = "—"

    p = doc.add_paragraph()
    p.add_run("Data și ora desfășurării: ").bold = True
    p.add_run(date_fmt)

    loc = meeting.get("loc_desfasurare") or meeting.get("locality") or "—"
    p = doc.add_paragraph()
    p.add_run("Locul desfășurării: ").bold = True
    p.add_run(str(loc))

    # Participanți
    vc = meeting.get("vertical_config") or {}
    participants = (
        vc.get("participanti")
        or vc.get("participants")
        or meeting.get("participanti")
        or []
    )
    if participants:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Participanți: ").bold = True
        if isinstance(participants, list):
            for item in participants:
                doc.add_paragraph(f"• {item}")
        else:
            p.add_run(str(participants))

    # ---- ORDINEA DE ZI ----
    agenda = (
        vc.get("ordine_de_zi")
        or vc.get("subiecte_discutate")
        or vc.get("agenda")
        or []
    )
    if agenda:
        doc.add_paragraph()
        h = doc.add_heading("ORDINEA DE ZI", level=2)
        if isinstance(agenda, list):
            for i, item in enumerate(agenda, 1):
                doc.add_paragraph(f"{i}. {item}")
        else:
            doc.add_paragraph(str(agenda))

    # ---- DEZBATERI (from diarized transcript) ----
    diarized = meeting.get("diarized_transcript") or []
    if diarized:
        doc.add_paragraph()
        doc.add_heading("DEZBATERI", level=2)
        for entry in diarized:
            speaker = entry.get("speaker") or "Vorbitor"
            role = entry.get("role") or ""
            timestamp = entry.get("timestamp") or ""
            text = entry.get("text") or ""
            p = doc.add_paragraph()
            label = f"{speaker}"
            if role:
                label += f" ({role})"
            if timestamp:
                label += f" — {timestamp}"
            label += ":"
            run = p.add_run(label)
            run.bold = True
            p.add_run(f" {text}")
    else:
        # Fallback: raw transcript
        transcript = meeting.get("transcript", "")
        if transcript:
            doc.add_paragraph()
            doc.add_heading("DEZBATERI", level=2)
            doc.add_paragraph(transcript)

    # ---- HOTĂRÂRI / DECIZII ----
    decisions = (
        vc.get("decizii")
        or vc.get("hotarari")
        or vc.get("decizii_luate")
        or []
    )
    if decisions:
        doc.add_paragraph()
        doc.add_heading("HOTĂRÂRI ADOPTATE", level=2)
        if isinstance(decisions, list):
            for i, item in enumerate(decisions, 1):
                doc.add_paragraph(f"{i}. {item}")
        else:
            doc.add_paragraph(str(decisions))

    # ---- ACȚIUNI DE URMAT ----
    actions = (
        vc.get("actiuni_de_urmat")
        or vc.get("actions")
        or meeting.get("actions")
        or []
    )
    if actions:
        doc.add_paragraph()
        doc.add_heading("ACȚIUNI STABILITE", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Acțiune'
        hdr[1].text = 'Responsabil'
        hdr[2].text = 'Termen'
        if isinstance(actions, list):
            for a in actions:
                row = table.add_row().cells
                if isinstance(a, dict):
                    row[0].text = str(a.get('text') or a.get('actiune') or '')
                    row[1].text = str(a.get('owner') or a.get('responsabil') or '-')
                    row[2].text = str(a.get('deadline') or a.get('termen') or '-')
                else:
                    row[0].text = str(a)
                    row[1].text = '-'
                    row[2].text = '-'

    # ---- CONCLUZII ----
    conclusions = vc.get("concluzii") or meeting.get("concluzia") or ""
    if conclusions:
        doc.add_paragraph()
        doc.add_heading("CONCLUZII", level=2)
        if isinstance(conclusions, list):
            for c in conclusions:
                doc.add_paragraph(str(c))
        else:
            doc.add_paragraph(str(conclusions))

    # ---- SEMNĂTURI ----
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Drept pentru care am încheiat prezentul proces-verbal.").italic = True

    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Președinte de ședință"
    table.rows[0].cells[1].text = "Secretar"
    table.rows[1].cells[0].text = "\n\n_______________________"
    table.rows[1].cells[1].text = "\n\n_______________________"

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Document generat automat cu Meetings.ro — {datetime.now(timezone.utc).strftime('%d.%m.%Y')}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Output
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_title = re.sub(r'[^\w\s-]', '', transliterate_ro(title_text)).strip().replace(' ', '_')[:50]
    filename = f"ProcesVerbal_{safe_title}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


def transliterate_ro(text: str) -> str:
    """Transliterate Romanian special characters for PDF compatibility."""
    if not text:
        return ""
    mapping = {
        'ă': 'a', 'Ă': 'A',
        'â': 'a', 'Â': 'A',
        'î': 'i', 'Î': 'I',
        'ș': 's', 'Ș': 'S',
        'ş': 's', 'Ş': 'S',
        'ț': 't', 'Ț': 'T',
        'ţ': 't', 'Ţ': 'T',
    }
    result = text
    for ro_char, latin_char in mapping.items():
        result = result.replace(ro_char, latin_char)
    return result


# ---- MEETING STATUS POLLING ----

@app.get("/api/meetings/{meeting_id}/status")
async def get_meeting_status(meeting_id: str, user: dict = Depends(get_current_user)):
    """Quick status check for polling."""
    meeting = await verify_meeting_ownership(meeting_id, user)
    return serialize_doc({"_id": meeting["_id"], "status": meeting.get("status"), "error": meeting.get("error"), "title": meeting.get("title"), "locality": meeting.get("locality")})


# ---- HEALTH CHECK ----

@app.get("/api/health")
async def health_check():
    """API health check."""
    return {"status": "ok", "service": "Meetings.ro API"}


# ==================== STRIPE PAYMENTS ====================

class CheckoutRequest(BaseModel):
    plan: str  # "pro" or "enterprise"
    interval: str  # "monthly" or "yearly"
    user_email: str


@app.post("/api/payments/create-checkout-session")
async def create_checkout_session(req: CheckoutRequest, user: dict = Depends(get_current_user)):
    """Create a Stripe Checkout Session for subscription."""
    price_key = f"{req.plan}_{req.interval}"
    price_id = STRIPE_PRICES.get(price_key)
    if not price_id:
        raise HTTPException(status_code=400, detail="Plan invalid")

    try:
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            mode="subscription",
            customer_email=req.user_email,
            line_items=[{"price": price_id, "quantity": 1}],
            success_url="https://meetings-ro-api.onrender.com/payment-success?session_id={CHECKOUT_SESSION_ID}",
            cancel_url="https://meetings-ro-api.onrender.com/payment-cancel",
        )
        return {"url": session.url}
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/payments/webhook")
async def stripe_webhook(request: Request):
    """Handle Stripe webhook events."""
    print(f"[Stripe] Webhook received. Secret configured: {bool(STRIPE_WEBHOOK_SECRET)}, Key configured: {bool(stripe.api_key)}")
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")

    if not STRIPE_WEBHOOK_SECRET:
        raise HTTPException(status_code=500, detail="Webhook secret not configured")

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, STRIPE_WEBHOOK_SECRET)
    except ValueError as e:
        print(f"[Stripe] Invalid payload: {e}")
        raise HTTPException(status_code=400, detail="Invalid payload")
    except stripe.error.SignatureVerificationError as e:
        print(f"[Stripe] Invalid signature: {e}")
        raise HTTPException(status_code=400, detail="Invalid signature")
    except Exception as e:
        print(f"[Stripe] UNHANDLED construct_event error: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=f"Webhook error: {str(e)}")

    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        user_email = session.get("customer_email")
        # Determine plan from price ID
        subscription_id = session.get("subscription")
        plan_tier = "PRO"  # default

        if subscription_id:
            try:
                sub = stripe.Subscription.retrieve(subscription_id)
                price_id = sub["items"]["data"][0]["price"]["id"]
                if price_id in [STRIPE_PRICES["enterprise_monthly"], STRIPE_PRICES["enterprise_yearly"]]:
                    plan_tier = "ENTERPRISE"
            except Exception:
                pass

        if user_email:
            await db.users.update_one(
                {"email": user_email},
                {"$set": {
                    "plan": plan_tier,
                    "stripe_customer_id": session.get("customer"),
                    "stripe_subscription_id": subscription_id,
                    "plan_updated_at": datetime.now(timezone.utc),
                }},
                upsert=True,
            )
            print(f"[Stripe] User {user_email} upgraded to {plan_tier}")

            # Auto-generate SmartBill invoice + e-Factura (non-blocking)
            issue_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
            asyncio.create_task(_create_and_save_invoice(
                user_email=user_email,
                plan=plan_tier,
                issue_date=issue_date,
            ))

    elif event["type"] == "customer.subscription.deleted":
        session = event["data"]["object"]
        customer_id = session.get("customer")
        if customer_id:
            await db.users.update_one(
                {"stripe_customer_id": customer_id},
                {"$set": {
                    "plan": "FREE",
                    "stripe_subscription_id": None,
                    "plan_updated_at": datetime.now(timezone.utc),
                }}
            )
            print(f"[Stripe] Customer {customer_id} downgraded to FREE")

    return {"status": "ok"}


@app.get("/payment-success")
async def payment_success(session_id: str = ""):
    """Redirect page after successful payment."""
    return JSONResponse(content={
        "status": "success",
        "message": "Plata a fost procesată cu succes! Poți închide această pagină și reveni în aplicație.",
        "session_id": session_id,
    })


@app.get("/payment-cancel")
async def payment_cancel():
    """Redirect page after cancelled payment."""
    return JSONResponse(content={
        "status": "cancelled",
        "message": "Plata a fost anulată. Poți reveni în aplicație.",
    })


# ==================== LEGAL PAGES (App Store / Google Play) ====================

@app.get("/privacy", response_class=HTMLResponse)
async def privacy_policy():
    return HTMLResponse("""
    <html>
    <head><title>Politică de Confidențialitate — Meetings.ro</title>
    <meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
    <style>body{font-family:'Segoe UI',system-ui,sans-serif;max-width:800px;margin:40px auto;padding:0 20px;line-height:1.7;color:#333}h1{color:#1B2A4A}h2{color:#1B2A4A;margin-top:28px}</style>
    </head>
    <body>
    <h1>Politică de Confidențialitate</h1>
    <p>Ultima actualizare: Aprilie 2026</p>
    <h2>Date colectate</h2>
    <p>Meetings.ro colectează: adresă email, numele utilizatorului, înregistrări audio ale ședințelor, transcrieri și rapoarte generate automat.</p>
    <h2>Utilizarea datelor</h2>
    <p>Datele sunt folosite exclusiv pentru generarea transcrierilor și rapoartelor solicitate de utilizator. Nu vindem și nu partajăm datele cu terți.</p>
    <h2>Stocarea datelor</h2>
    <p>Datele sunt stocate pe servere securizate în Europa (Frankfurt, Germania) prin intermediul platformei Render.com și MongoDB Atlas.</p>
    <h2>Servicii terțe</h2>
    <p>Folosim: Groq/OpenAI (transcriere audio), Anthropic Claude (extracție date), Stripe (plăți), Resend (email-uri tranzacționale). Fiecare serviciu procesează doar datele minime necesare.</p>
    <h2>Drepturi GDPR</h2>
    <p>Ai dreptul la acces, rectificare, portabilitate și ștergere a datelor tale. Pentru orice solicitare, contactează-ne la adresa de mai jos.</p>
    <h2>Retenția datelor</h2>
    <p>Înregistrările audio și transcrierile sunt păstrate atât timp cât contul este activ. La ștergerea contului, toate datele sunt eliminate în maxim 30 de zile.</p>
    <h2>Contact</h2>
    <p>hello@meetings.ro</p>
    </body></html>
    """)


@app.get("/terms", response_class=HTMLResponse)
async def terms_of_service():
    return HTMLResponse("""
    <html>
    <head><title>Termeni și Condiții — Meetings.ro</title>
    <meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
    <style>body{font-family:'Segoe UI',system-ui,sans-serif;max-width:800px;margin:40px auto;padding:0 20px;line-height:1.7;color:#333}h1{color:#1B2A4A}h2{color:#1B2A4A;margin-top:28px}</style>
    </head>
    <body>
    <h1>Termeni și Condiții</h1>
    <p>Ultima actualizare: Aprilie 2026</p>
    <h2>Serviciul</h2>
    <p>Meetings.ro oferă servicii de transcriere automată și generare rapoarte AI pentru ședințe profesionale, disponibile prin aplicație mobilă.</p>
    <h2>Cont și responsabilitate</h2>
    <p>Ești responsabil pentru securitatea contului tău și pentru conținutul înregistrărilor uploadate. Nu este permisă înregistrarea fără consimțământul participanților.</p>
    <h2>Planuri și plăți</h2>
    <p>Planul gratuit include 5 întâlniri pe lună. Abonamentele plătite (Pro, Enterprise) sunt lunare sau anuale, procesate prin Stripe. Anularea se poate face oricând, cu efect la finalul perioadei plătite.</p>
    <h2>Proprietate intelectuală</h2>
    <p>Conținutul înregistrărilor și transcrierilor tale îți aparține. Meetings.ro deține drepturile asupra platformei, designului și algoritmilor.</p>
    <h2>Limitarea răspunderii</h2>
    <p>Meetings.ro nu garantează acuratețea 100% a transcrierilor sau a rapoartelor generate. Serviciul este furnizat "așa cum este".</p>
    <h2>Modificări ale termenilor</h2>
    <p>Ne rezervăm dreptul de a modifica acești termeni. Utilizatorii vor fi notificați prin email cu minim 14 zile înainte.</p>
    <h2>Contact</h2>
    <p>hello@meetings.ro</p>
    </body></html>
    """)

