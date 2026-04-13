"""
Export router — PDF and DOCX generation.
Diacritics bug FIXED: uses DejaVu Unicode font instead of stripping chars.
Roles: anyone with 'export' permission (mayor, councilor, secretary, admin).
"""
import re
from datetime import datetime
from io import BytesIO
from typing import Optional

from bson import ObjectId
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse

from auth.rbac import get_current_user, require_permission
from db.collections import meetings_col

router = APIRouter(prefix="/api/meetings", tags=["export"])

DEJAVU_FONT_PATH = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
DEJAVU_BOLD_PATH = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"


# ── Helpers ────────────────────────────────────────────────────────────────────

def serialize_doc(doc):
    if doc is None:
        return None
    result = {}
    for key, value in doc.items():
        if isinstance(value, ObjectId):
            result[key] = str(value)
        elif isinstance(value, datetime):
            result[key] = value.isoformat()
        elif isinstance(value, list):
            result[key] = [
                serialize_doc(v) if isinstance(v, dict)
                else (str(v) if isinstance(v, ObjectId) else v)
                for v in value
            ]
        elif isinstance(value, dict):
            result[key] = serialize_doc(value)
        else:
            result[key] = value
    return result


async def get_meeting_for_export(meeting_id: str, user: dict) -> dict:
    """Fetch meeting with tenant isolation check."""
    try:
        meeting = await meetings_col().find_one({"_id": ObjectId(meeting_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="ID meeting invalid")
    if not meeting:
        raise HTTPException(status_code=404, detail="Ședința nu a fost găsită")
    if str(meeting.get("tenant_id", "")) != str(user.get("tenant_id", "")):
        raise HTTPException(status_code=403, detail="Nu ai acces la această ședință")
    return meeting


def _format_date(date_str: Optional[str]) -> str:
    if not date_str:
        return ""
    try:
        dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        return dt.strftime("%d.%m.%Y %H:%M")
    except (ValueError, AttributeError):
        return date_str


def _safe_filename(text: str, max_len: int = 40) -> str:
    return re.sub(r"[^\w\s-]", "", text).strip().replace(" ", "_")[:max_len]


# ── PDF export (with Unicode diacritics support) ───────────────────────────────

@router.get("/{meeting_id}/export/pdf")
async def export_pdf(
    meeting_id: str,
    user: dict = Depends(require_permission("export")),
):
    """Export meeting as PDF. Supports full Romanian diacritics via DejaVu font."""
    from fpdf import FPDF

    meeting = await get_meeting_for_export(meeting_id, user)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)

    # Add Unicode font — fixes the diacritics bug in the original transliterate_ro()
    import os
    if os.path.exists(DEJAVU_FONT_PATH):
        pdf.add_font("DejaVu", "", DEJAVU_FONT_PATH, uni=True)
        pdf.add_font("DejaVu", "B", DEJAVU_BOLD_PATH, uni=True)
        regular_font = "DejaVu"
        bold_font = "DejaVu"
    else:
        # Fallback: Arial (no diacritics, but won't crash)
        regular_font = "Arial"
        bold_font = "Arial"

    title = meeting.get("title") or "Ședință"
    locality = meeting.get("locality") or "Necunoscut"
    date_str = _format_date(meeting.get("date", ""))
    from config import settings
    institution = settings.INSTITUTION_NAME

    # Header
    pdf.set_font(bold_font, "B", 14)
    pdf.cell(0, 8, institution, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(bold_font, "B", 16)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(regular_font, "", 10)
    pdf.cell(0, 6, f"Localitate: {locality}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Data: {date_str}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Vertical config fields
    vertical_config = meeting.get("vertical_config", {})
    if vertical_config:
        pdf.set_font(bold_font, "B", 12)
        pdf.cell(0, 8, "Raport", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(regular_font, "", 9)
        for k, v in vertical_config.items():
            if v and k != "locality":
                label = k.replace("_", " ").title()
                value_str = str(v)[:500]
                pdf.set_font(bold_font, "B", 9)
                pdf.cell(0, 5, f"{label}:", new_x="LMARGIN", new_y="NEXT")
                pdf.set_font(regular_font, "", 9)
                pdf.multi_cell(0, 5, value_str)
                pdf.ln(1)
        pdf.ln(3)

    # Summary
    summary = meeting.get("summary", [])
    if summary:
        pdf.set_font(bold_font, "B", 12)
        pdf.cell(0, 8, "Rezumat", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(regular_font, "", 9)
        for i, item in enumerate(summary[:10]):
            pdf.cell(0, 5, f"{i+1}. {str(item)[:120]}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

    # Action items
    actions = meeting.get("actions", [])
    if actions:
        pdf.set_font(bold_font, "B", 12)
        pdf.cell(0, 8, "Acțiuni", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(regular_font, "", 9)
        for action in actions[:20]:
            status_mark = "[X]" if action.get("completed") else "[ ]"
            text = str(action.get("text", ""))[:100]
            pdf.cell(0, 5, f"{status_mark} {text}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

    # Transcript (truncated for PDF)
    transcript = meeting.get("transcript", "")
    if transcript:
        pdf.set_font(bold_font, "B", 12)
        pdf.cell(0, 8, "Transcriere", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(regular_font, "", 8)
        pdf.multi_cell(0, 4, transcript[:3000])

    pdf_bytes = pdf.output()
    buffer = BytesIO(pdf_bytes)
    safe_title = _safe_filename(title)
    safe_date = date_str.replace(".", "-").replace(" ", "_").replace(":", "-") if date_str else "no-date"
    filename = f"sedinta_{safe_title}_{safe_date}.pdf"

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── DOCX export ────────────────────────────────────────────────────────────────

@router.get("/{meeting_id}/export/docx")
async def export_docx(
    meeting_id: str,
    user: dict = Depends(require_permission("export")),
):
    """Export meeting as DOCX. Full Romanian diacritics supported natively."""
    from docx import Document
    from docx.shared import Pt
    from config import settings

    meeting = await get_meeting_for_export(meeting_id, user)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = meeting.get("title") or "Ședință"
    locality = meeting.get("locality") or "Necunoscut"
    date_str = _format_date(meeting.get("date", ""))

    doc.add_heading(settings.INSTITUTION_NAME, level=2)
    doc.add_heading(title, level=1)

    p = doc.add_paragraph()
    p.add_run("Localitate: ").bold = True
    p.add_run(locality)
    p = doc.add_paragraph()
    p.add_run("Data: ").bold = True
    p.add_run(date_str)

    # Vertical config fields
    vertical_config = meeting.get("vertical_config", {})
    if vertical_config:
        doc.add_heading("Raport", level=2)
        for k, v in vertical_config.items():
            if v and k != "locality":
                label = k.replace("_", " ").title()
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(v))

    # Summary
    summary = meeting.get("summary", [])
    if summary:
        doc.add_heading("Rezumat", level=2)
        for item in summary:
            doc.add_paragraph(str(item), style="List Bullet")

    # Key Points
    key_points = meeting.get("key_points", [])
    if key_points:
        doc.add_heading("Puncte cheie", level=2)
        for item in key_points:
            doc.add_paragraph(str(item), style="List Bullet")

    # Action items table
    actions = meeting.get("actions", [])
    if actions:
        doc.add_heading("Acțiuni", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Status", "Acțiune", "Responsabil", "Termen"
        for action in actions:
            row = table.add_row().cells
            row[0].text = "✓" if action.get("completed") else "○"
            row[1].text = action.get("text", "")
            row[2].text = action.get("owner", "-") or "-"
            row[3].text = action.get("deadline", "-") or "-"

    # Transcript
    transcript = meeting.get("transcript", "")
    if transcript:
        doc.add_heading("Transcriere", level=2)
        doc.add_paragraph(transcript)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    safe_title = _safe_filename(title)
    filename = f"sedinta_{safe_title}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
