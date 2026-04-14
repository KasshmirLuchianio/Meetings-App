"""
Romanian institutional document export — "Proces Verbal".

Standard format required by Romanian public authorities (UAT, consilii locale):
  - Font: Times New Roman 12pt
  - Line spacing: 1.5
  - Page: A4, 2.5cm margins all sides
  - Sections: Participanți, Ordinea de zi, Dezbateri, Hotărâri
  - Footer: Secretar / Președinte de ședință signature blocks

Generates DOCX natively (full diacritics via python-docx).
Converts to PDF via LibreOffice --headless (if installed).
"""
import io
import os
import subprocess
import tempfile
from datetime import datetime
from typing import Optional


# ── DOCX generation ─────────────────────────────────────────────────────────────

def generate_docx(meeting: dict, institution_name: str) -> io.BytesIO:
    """
    Build a complete 'Proces Verbal' DOCX from a meeting document.

    Reads from meeting['proces_verbal'] if a structured report exists.
    Falls back to meeting['diarized_segments'] → plain transcript as needed.

    Returns a BytesIO buffer ready to stream.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    doc = Document()

    # ── Page layout: A4, 2.5cm margins ─────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width  = Cm(21.0)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

    # ── Default style: Times New Roman 12pt, 1.5 line spacing ──────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    _set_spacing(normal.paragraph_format, 1.5)

    # Heading styles
    for level, pt_size in [(1, 14), (2, 13), (3, 12)]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Times New Roman"
        h_style.font.size = Pt(pt_size)
        h_style.font.bold = True
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        _set_spacing(h_style.paragraph_format, 1.5)

    # ── Unpack structured report or fall back to raw data ──────────────────────
    report      = meeting.get("proces_verbal") or {}
    antet       = report.get("antet") or {}
    participanti = report.get("participanti") or []
    ordine_de_zi = report.get("ordine_de_zi") or []
    dezbateri    = report.get("dezbateri") or []
    hotarari     = report.get("hotarari") or []
    observatii   = report.get("observatii")

    locality = (
        meeting.get("locality")
        or antet.get("locatie")
        or "Necunoscut"
    )
    date_str  = _format_date_ro(meeting.get("date") or meeting.get("created_at", ""))

    # ════════════════════════════════════════════════════════════════════════════
    # HEADER BLOCK
    # ════════════════════════════════════════════════════════════════════════════

    # Logo placeholder
    _add_centered_run(doc, "[LOGO INSTITUȚIE]", Pt(9), RGBColor(180, 180, 180), italic=True)

    # Institution name — centered, caps, bold
    _add_centered_run(doc, institution_name.upper(), Pt(13), bold=True)

    # "PROCES VERBAL" title
    _add_centered_run(doc, "PROCES VERBAL", Pt(15), bold=True)

    # Session type (if available)
    tip = antet.get("tip_sedinta", "")
    if tip:
        _add_centered_run(doc, f"al {tip}", Pt(12), italic=True)

    # Date / location line
    ora_start = antet.get("ora_inceput", "")
    ora_end   = antet.get("ora_sfarsit", "")
    ora_str = f", ora {ora_start}" if ora_start else ""
    if ora_start and ora_end:
        ora_str = f", orele {ora_start}—{ora_end}"
    locatie = antet.get("locatie") or locality
    _add_centered_run(doc, f"{date_str}{ora_str}, {locatie}", Pt(12))

    doc.add_paragraph()  # spacer

    # ════════════════════════════════════════════════════════════════════════════
    # 1. PARTICIPANȚI
    # ════════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "1. PARTICIPANȚI")

    if participanti:
        for part in participanti:
            name = part.get("nume") or part.get("name") or ""
            role = part.get("functie") or part.get("role") or ""
            p = _body_para(doc)
            r = p.add_run(f"{name}")
            r.bold = True
            if role:
                p.add_run(f" — {role}")
    else:
        # Build from diarized segments
        diarized = meeting.get("diarized_segments", [])
        seen: dict = {}
        for seg in diarized:
            spk = seg.get("speaker", "")
            if spk and spk not in seen:
                seen[spk] = seg.get("role", "")
        if seen:
            for name, role in seen.items():
                p = _body_para(doc)
                p.add_run(name).bold = True
                if role:
                    p.add_run(f" — {role}")
        else:
            np_ = meeting.get("numar_participanti")
            _body_para(doc).add_run(f"Număr participanți: {np_ or 'Nespecificat'}")

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════════
    # 2. ORDINEA DE ZI
    # ════════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "2. ORDINEA DE ZI")

    if ordine_de_zi:
        for i, item in enumerate(ordine_de_zi, 1):
            p = _body_para(doc)
            p.add_run(f"{i}. {item}")
    else:
        _body_para(doc).add_run("—")

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════════
    # 3. DEZBATERI
    # ════════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "3. DEZBATERI")

    if dezbateri:
        for punct in dezbateri:
            punct_nr    = punct.get("punct", "")
            punct_titlu = punct.get("titlu", "")

            # Agenda-point sub-heading
            p = _body_para(doc)
            label = f"Punctul {punct_nr}. {punct_titlu}" if punct_nr else punct_titlu
            run = p.add_run(label)
            run.bold = True
            run.underline = True

            for interv in (punct.get("interventii") or []):
                speaker   = interv.get("speaker", "")
                functie   = interv.get("functie", "")
                continut  = interv.get("continut", "")
                timestamp = interv.get("timestamp", "")

                p_int = _body_para(doc, left_indent_cm=1.0)
                spk_label = speaker
                if functie:
                    spk_label += f" ({functie})"
                if timestamp:
                    spk_label += f" [{timestamp}]"
                p_int.add_run(f"{spk_label}: ").bold = True
                p_int.add_run(continut)

    else:
        # Fallback to diarized segments
        diarized = meeting.get("diarized_segments", [])
        if diarized:
            for seg in diarized:
                speaker   = seg.get("speaker", "Vorbitor")
                role      = seg.get("role", "")
                text      = seg.get("text", "")
                timestamp = seg.get("timestamp", "")
                if not text.strip():
                    continue
                p = _body_para(doc, left_indent_cm=1.0)
                label = speaker
                if role:
                    label += f" ({role})"
                if timestamp:
                    label += f" [{timestamp}]"
                p.add_run(f"{label}: ").bold = True
                p.add_run(text)
        elif meeting.get("transcript"):
            _body_para(doc).add_run(meeting["transcript"])

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════════
    # 4. HOTĂRÂRI
    # ════════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "4. HOTĂRÂRI")

    if hotarari:
        for i, hot in enumerate(hotarari, 1):
            p = _body_para(doc)
            p.add_run(f"{i}. ").bold = True
            p.add_run(str(hot))
    else:
        _body_para(doc).add_run("—")

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════════
    # 5. OBSERVAȚII (optional)
    # ════════════════════════════════════════════════════════════════════════════
    if observatii:
        _section_heading(doc, "5. OBSERVAȚII")
        _body_para(doc).add_run(str(observatii))
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════════
    # SIGNATURE BLOCK
    # ════════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    _add_signature_block(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── PDF via LibreOffice ─────────────────────────────────────────────────────────

def generate_pdf_from_docx(docx_buf: io.BytesIO) -> Optional[io.BytesIO]:
    """
    Convert a DOCX buffer to PDF using LibreOffice --headless.
    Returns None if LibreOffice is not installed or conversion fails.
    """
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "document.docx")
            pdf_path  = os.path.join(tmpdir, "document.pdf")

            with open(docx_path, "wb") as f:
                f.write(docx_buf.getvalue())

            result = subprocess.run(
                [
                    "libreoffice", "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    docx_path,
                ],
                capture_output=True,
                timeout=60,
            )

            if result.returncode != 0:
                print(f"[ProcesVerbal] LibreOffice error: {result.stderr.decode()}")
                return None

            if not os.path.exists(pdf_path):
                return None

            with open(pdf_path, "rb") as f:
                return io.BytesIO(f.read())

    except FileNotFoundError:
        print("[ProcesVerbal] LibreOffice not found — PDF conversion unavailable.")
        return None
    except Exception as exc:
        print(f"[ProcesVerbal] PDF conversion failed: {exc}")
        return None


# ── Document helpers ─────────────────────────────────────────────────────────────

def _set_spacing(pf, multiple: float):
    from docx.enum.text import WD_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = multiple


def _body_para(doc, left_indent_cm: float = 0.0):
    """Add a normal body paragraph with 1.5 line spacing."""
    from docx.shared import Cm
    p = doc.add_paragraph()
    _set_spacing(p.paragraph_format, 1.5)
    if left_indent_cm:
        p.paragraph_format.left_indent = Cm(left_indent_cm)
    return p


def _section_heading(doc, text: str):
    """Add a bold, underlined section heading."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    _set_spacing(p.paragraph_format, 1.5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = __import__("docx").shared.Pt(12)


def _add_centered_run(doc, text: str, size, color=None, bold=False, italic=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p.paragraph_format, 1.5)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def _add_signature_block(doc):
    """
    Two-column signature block:
        SECRETAR,              PREȘEDINTE DE ȘEDINȚĂ,
        _________________________   _________________________
    """
    from docx.shared import Pt, Cm

    table = doc.add_table(rows=2, cols=2)
    # Remove all borders (invisible layout table)
    for row in table.rows:
        for cell in row.cells:
            _remove_borders(cell)

    labels = ["SECRETAR,", "PREȘEDINTE DE ȘEDINȚĂ,"]
    for i, label in enumerate(labels):
        cell = table.rows[0].cells[i]
        para = cell.paragraphs[0]
        _set_spacing(para.paragraph_format, 1.5)
        run = para.add_run(label)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    for i in range(2):
        cell = table.rows[1].cells[i]
        para = cell.paragraphs[0]
        _set_spacing(para.paragraph_format, 1.5)
        run = para.add_run("_" * 32)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _remove_borders(cell):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "none")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _format_date_ro(date_str: Optional[str]) -> str:
    """Format an ISO date string to Romanian long format: '14 aprilie 2026'."""
    if not date_str:
        return ""
    try:
        dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        months = [
            "ianuarie", "februarie", "martie", "aprilie", "mai", "iunie",
            "iulie", "august", "septembrie", "octombrie", "noiembrie", "decembrie",
        ]
        return f"{dt.day} {months[dt.month - 1]} {dt.year}"
    except (ValueError, AttributeError):
        return date_str
